# ---------------------------------------------------------
# IMPORTACIONES DEL SISTEMA
# ---------------------------------------------------------
import streamlit as st
import pandas as pd
import os
import json
import uuid
import gspread
from google.oauth2.service_account import Credentials
from datetime import datetime, timedelta

# ---------------------------------------------------------
# BLOQUE 0: CONFIGURACIÓN GENERAL Y PERSISTENCIA
# ---------------------------------------------------------
import streamlit as st
import pandas as pd
import os
import json
import uuid
from datetime import datetime, timedelta
import gspread
from google.oauth2.service_account import Credentials

st.set_page_config(page_title="JPV - OpsControl", layout="wide")

PERSISTENCE_DIR = "persistence"

def init_system():
    os.makedirs(PERSISTENCE_DIR, exist_ok=True)

def get_week_identifier(offset_weeks=0):
    target_date = datetime.now() + timedelta(weeks=offset_weeks)
    return target_date.strftime("%Y_W%W")

def apply_custom_styles():
    st.markdown("""
        <style>
        .main { background-color: #f5f7f9; }
        .stButton>button { width: 100%; border-radius: 5px; height: 3em; background-color: #004a99; color: white; font-weight: bold;}
        .btn-guardar>button { background-color: #217346; color: white; width: 100%; height: 3em; margin-top: 20px;}
        .stDataFrame { border: 1px solid #c4ced4; }
        .marco-caso { background-color: white; padding: 15px; border-radius: 5px; border-left: 5px solid #217346; margin-bottom: 10px; box-shadow: 0 1px 3px rgba(0,0,0,0.1); }
        .marco-gestion { background-color: white; padding: 15px; border-radius: 5px; border-left: 5px solid #004a99; margin-bottom: 10px; box-shadow: 0 1px 3px rgba(0,0,0,0.1); }
        .tarea-marco { background-color: white; padding: 15px; border-radius: 5px; border-left: 5px solid #004a99; margin-bottom: 10px; box-shadow: 0 1px 3px rgba(0,0,0,0.1); }
        .tarea-realizada { border-left: 5px solid #217346; opacity: 0.8; }
        h1, h2, h3, h4 { color: #003366 !important; font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; }
        </style>
    """, unsafe_allow_html=True)

init_system()
apply_custom_styles()

# ---------------------------------------------------------
# BLOQUE 1: FUNCIONES DE MEMORIA Y BASE DE DATOS LOCAL/NUBE
# VERSIÓN: 2.1.5 (Advertencia de Filtros en Reporte de Acciones)
# ---------------------------------------------------------
def get_google_client():
    try:
        if "gcp_service_account" in st.secrets and "google_sheet_url" in st.secrets:
            scope = ['https://www.googleapis.com/auth/spreadsheets', 'https://www.googleapis.com/auth/drive']
            creds = Credentials.from_service_account_info(st.secrets["gcp_service_account"], scopes=scope)
            return gspread.authorize(creds)
    except Exception as e:
        st.error(f"⚠️ Error crítico de conexión a Google Cloud: {e}")
    return None

def get_google_sheet():
    client = get_google_client()
    if client:
        try:
            return client.open_by_url(st.secrets["google_sheet_url"]).sheet1
        except Exception as e:
            st.error(f"⚠️ Error al abrir la hoja de cálculo: {e}")
    return None

def load_master_base():
    st.sidebar.header("📁 Base Maestra")
    filepath = os.path.join(PERSISTENCE_DIR, "BASE_MAESTRA.json")
    
    df_local = None
    fecha_actualizacion = None
    
    # 1. Recuperación en silencio desde Google Sheets si no hay base local
    if not os.path.exists(filepath):
        client = get_google_client()
        if client:
            try:
                doc = client.open_by_url(st.secrets["google_sheet_url"])
                ws = doc.worksheet("Base_Maestra")
                metadata = ws.row_values(1)
                
                if len(metadata) >= 2 and metadata[0] == "FECHA_ACTUALIZACION":
                    fecha_str = metadata[1]
                    datos_crud = ws.get_all_records(head=2) 
                    df_local = pd.DataFrame(datos_crud)
                    
                    datos_guardar = {"fecha": fecha_str, "data": df_local.to_dict(orient="records")}
                    with open(filepath, 'w', encoding='utf-8') as f:
                        json.dump(datos_guardar, f, ensure_ascii=False)
            except Exception:
                pass
                
    # 2. Lectura de caché local
    if os.path.exists(filepath):
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                datos = json.load(f)
                df_local = pd.DataFrame(datos['data'])
                fecha_actualizacion = datetime.strptime(datos['fecha'], "%Y-%m-%d %H:%M:%S")
        except Exception:
            pass

    # 3. Lógica de Semáforo y Caducidad (7 Días)
    necesita_actualizacion = True
    if fecha_actualizacion:
        dias = (datetime.now() - fecha_actualizacion).days
        if dias <= 7:
            st.sidebar.success(f"✅ Base actualizada hace {dias} días ({fecha_actualizacion.strftime('%d/%m/%Y')})")
            necesita_actualizacion = False
        else:
            st.sidebar.error(f"🚨 Base caducada (hace {dias} días). Requiere actualización urgente.")
    else:
        st.sidebar.warning("⚠️ No hay base de datos en el sistema.")

    # 4. Motor de Carga, Sanitización y Respaldo Nube
    with st.sidebar.expander("📥 Subir / Actualizar Base Maestra", expanded=necesita_actualizacion):
        
        # --- ADVERTENCIA ESTRATÉGICA ANTES DE CARGAR ---
        st.warning("💡 **Requisito del Excel:** El reporte de acciones extraído del sistema debe contemplar a **todas las divisiones de la gerencia**. Antes de subirlo, verifique que no existan filtros que oculten a los ajustadores y asegúrese de **excluir** los casos en estado *Anulado* o *Cerrado*.")
        
        uploaded_file = st.file_uploader("Cargar 'Reporte de Acciones'", type=["xlsx", "csv"])
        if uploaded_file is not None:
            # Creamos una huella digital única para el archivo subido
            file_signature = f"{uploaded_file.name}_{uploaded_file.size}"
            
            # El Guardián solo permite procesar si es un archivo nuevo o no ha sido registrado en esta sesión
            if st.session_state.get("ultima_base_procesada") != file_signature:
                try:
                    if uploaded_file.name.endswith('.xlsx'):
                        df_nuevo = pd.read_excel(uploaded_file, skiprows=5)
                    else:
                        df_nuevo = pd.read_csv(uploaded_file, skiprows=5)
                    
                    df_nuevo = df_nuevo.fillna("") 
                    for col in df_nuevo.columns:
                        df_nuevo[col] = df_nuevo[col].astype(str)
                        df_nuevo[col] = df_nuevo[col].apply(
                            lambda x: "" if str(x).strip().lower() in ["nan", "nat", "none", "<na>", "inf", "-inf"] else x
                        )

                    fecha_hoy = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    
                    # Guardado Local Inmediato (Asegura operatividad interna)
                    datos_guardar = {"fecha": fecha_hoy, "data": df_nuevo.to_dict(orient="records")}
                    with open(filepath, 'w', encoding='utf-8') as f:
                        json.dump(datos_guardar, f, ensure_ascii=False)
                        
                    # Intento de Sincronización en la Nube corporativa
                    try:
                        client = get_google_client()
                        if client:
                            doc = client.open_by_url(st.secrets["google_sheet_url"])
                            try:
                                ws = doc.worksheet("Base_Maestra")
                            except:
                                ws = doc.add_worksheet(title="Base_Maestra", rows="100", cols="100")
                            
                            ws.clear()
                            matriz = [["FECHA_ACTUALIZACION", fecha_hoy]]
                            matriz.append(df_nuevo.columns.astype(str).tolist())
                            matriz.extend(df_nuevo.values.tolist())
                            ws.update("A1", matriz)

                        st.success("¡Base Maestra procesada y asegurada en la nube corporativa!")
                    except Exception as cloud_error:
                        if "429" in str(cloud_error):
                            st.warning("⚠️ Base guardada localmente con éxito. Google Cloud está en pausa (Límite 429 de peticiones). Puedes operar tu planificador normalmente.")
                        else:
                            st.warning(f"⚠️ Base guardada localmente. Hubo un detalle con el respaldo en nube: {cloud_error}")
                    
                    # Registramos el archivo como procesado para congelar ejecuciones repetidas
                    st.session_state["ultima_base_procesada"] = file_signature
                    st.rerun() 
                except Exception as e:
                    st.error(f"Error crítico al procesar el Excel: {e}")
                
    return df_local

def limpiar_monto_mcl(valor):
    if pd.isna(valor) or str(valor).strip() in ["", "nan", "NaN", "NaT"]: return 0.0
    if isinstance(valor, (int, float)): return float(valor)
    
    v_str = str(valor).strip().replace('$', '').replace(' ', '')
    if '.' in v_str and ',' in v_str:
        if v_str.rfind(',') > v_str.rfind('.'):
            v_str = v_str.replace('.', '').replace(',', '.')
        else:
            v_str = v_str.replace(',', '')
    elif ',' in v_str:
        v_str = v_str.replace(',', '.')
        
    try:
        return float(v_str)
    except:
        return 0.0

def calcular_tramo_mcl(fila):
    valor = 0.0
    divisa = str(fila.get('Divisa', '')).upper()
    col_perdida = 'Perdida bruta (en moneda del caso)'
    
    if col_perdida in fila and pd.notna(fila[col_perdida]):
        valor = limpiar_monto_mcl(fila[col_perdida])

    is_mcl = False
    tramo_str = "<= 1000 UF"
    
    if 'USD' in divisa or 'US$' in divisa or 'DÓLAR' in divisa or 'DOLAR' in divisa:
        if valor > 200000:
            is_mcl = True
            tramo_str = "> 200.000 USD (MCL)"
        else:
            tramo_str = "<= 200.000 USD"
    else: 
        if valor <= 1000: tramo_str = "<= 1000 UF"
        elif valor <= 5000: tramo_str = "> 1000 Y <= 5000 UF"
        else:
            is_mcl = True
            tramo_str = "> 5000 UF (MCL)"
            
    return tramo_str, is_mcl

def get_month_identifier(offset_months=0):
    now = datetime.now()
    year = now.year
    month = now.month + offset_months
    while month > 12:
        month -= 12
        year += 1
    return f"{year}_{month:02d}"

def sync_from_cloud(filename, filepath):
    sheet = get_google_sheet()
    if sheet:
        try:
            registros = sheet.get_all_records()
            for fila in registros:
                if str(fila.get('Archivo', '')) == filename:
                    datos_json = json.loads(fila.get('JSON_Data', '[]'))
                    with open(filepath, 'w', encoding='utf-8') as f:
                        json.dump(datos_json, f, ensure_ascii=False, indent=4)
                    return datos_json
        except Exception as e:
            pass
    return []

def load_plan_semanal(ajustador, offset_weeks=0):
    week_id = get_week_identifier(offset_weeks)
    filename = f"plan_{ajustador.replace(' ', '_')}_{week_id}.json"
    filepath = os.path.join(PERSISTENCE_DIR, filename)
    if os.path.exists(filepath):
        with open(filepath, 'r', encoding='utf-8') as f:
            return json.load(f), filepath
    else:
        data = sync_from_cloud(filename, filepath)
        return data, filepath 

def load_plan_mensual(ajustador, offset_months=0, explicit_month_id=None):
    month_id = explicit_month_id if explicit_month_id else get_month_identifier(offset_months)
    filename = f"plan_mensual_mcl_{ajustador.replace(' ', '_')}_{month_id}.json"
    filepath = os.path.join(PERSISTENCE_DIR, filename)
    if os.path.exists(filepath):
        with open(filepath, 'r', encoding='utf-8') as f:
            return json.load(f), filepath
    else:
        data = sync_from_cloud(filename, filepath)
        return data, filepath

def save_plan_actualizado(filepath, data):
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=4)
        
    sheet = get_google_sheet()
    if sheet:
        try:
            filename = os.path.basename(filepath)
            json_str = json.dumps(data, ensure_ascii=False)
            fecha_upd = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            try:
                headers = sheet.row_values(1)
            except:
                headers = []
                
            if not headers or 'Archivo' not in headers:
                sheet.clear()
                sheet.append_row(['Archivo', 'JSON_Data', 'Ultima_Actualizacion'])
                sheet.append_row([filename, json_str, fecha_upd])
                return

            col_archivo_idx = headers.index('Archivo') + 1
            archivos_en_nube = sheet.col_values(col_archivo_idx)
            
            if filename in archivos_en_nube:
                fila_idx = archivos_en_nube.index(filename) + 1
                rango = f'A{fila_idx}:C{fila_idx}'
                sheet.update(rango, [[filename, json_str, fecha_upd]])
            else:
                sheet.append_row([filename, json_str, fecha_upd])
        except Exception as e:
            st.error(f"❌ Error al escribir filas en Google Sheets: {e}")

# ---------------------------------------------------------
# BLOQUE 2: VISTA - PLANIFICADOR (SEMANAL Y MENSUAL MCL)
# VERSIÓN: 2.4.4 (Candado Inteligente con Extensión por Feriados)
# ---------------------------------------------------------
def vista_planificador(modo="Semanal"):
    import pandas as pd
    import json
    import os
    import uuid
    import pytz
    from datetime import datetime, timedelta
    
    # Motor de feriados (Requiere añadir 'holidays' a requirements.txt)
    try:
        import holidays
        feriados_cl = holidays.Chile()
    except ImportError:
        feriados_cl = []

    # --- RELOJ Y ZONA HORARIA (CHILE) ---
    tz_chile = pytz.timezone('America/Santiago')
    ahora_chile = datetime.now(tz_chile)
    hoy_dt = ahora_chile.date()

    col_t1, col_t2 = st.columns([2, 1])
    with col_t1:
        if modo == "Semanal":
            st.title("🗓️ Planificador Semanal")
            st.markdown("Seleccione los casos de la Base Maestra que proyecta gestionar.")
        else:
            st.title("🏆 Planificador Mensual MCL")
            st.markdown("Gestión estratégica de casos complejos (Major and Complex Losses > 5000 UF / > 200k USD).")
            
    with col_t2:
        st.markdown("<br>", unsafe_allow_html=True)
        if modo == "Semanal":
            semana_opcion = st.radio("¿Qué semana estás planificando?", ["Semana Actual", "Próxima Semana"], horizontal=True)
            offset_weeks = 0 if semana_opcion == "Semana Actual" else 1
            offset_months = 0
        else:
            mes_opcion = st.radio("¿Qué mes estás planificando?", ["Mes Actual", "Próximo Mes"], horizontal=True)
            offset_months = 0 if mes_opcion == "Mes Actual" else 1
            offset_weeks = 0
            
        # Reloj visible para el ajustador
        st.markdown(f"<div style='text-align: right; color: #6c757d; font-size: 14px; margin-top: 5px;'>🕒 Hora Oficial del Sistema (Chile): <b>{ahora_chile.strftime('%H:%M')} hrs</b></div>", unsafe_allow_html=True)
            
    CATALOGO_ACCIONES = {
        "En Ajuste": ["Revisión de cobertura", "Revisión de antecedentes", "Otro / Manual"],
        "Inspección": ["Presencial", "Remota", "Otro / Manual"],
        "Correos": ["Solicitud de Antecedentes", "Reiteracion 1", "Reiteracion 2", "Reiteracion 3", "Ultimatum", "Cierre por falta de interés", "Otro / Manual"],
        "Reunión": ["Presencial", "Presentación pptx", "On line", "Presentacion on line", "Otro / Manual"],
        "Preparar Informe": ["Preliminar Extendido", "Preliminar Corto", "Carta de Análisis de Pérdidas", "Carta de Cobertura (Rechazo)", "Informe Intermedio 1", "Informe Intermedio 2", "Informe Intermedio 3", "Informe Intermedio 4", "Informe Intermedio 5", "Informe Intermedio", "Informe Final de Liquidación", "Respuesta a Impugnación", "Ademdum", "Otro / Manual"],
        "Otra Acción (Manual)": ["Describir manualmente"]
    }
    
    # --- CANDADO TEMPORAL INTELIGENTE (VENTANA DE COMPROMISO) ---
    es_adicional = False
    if modo == "Semanal":
        target_date = ahora_chile + timedelta(weeks=offset_weeks)
        lunes_target = target_date.date() - timedelta(days=target_date.weekday())
        viernes_prev = lunes_target - timedelta(days=3)
        
        # Lógica de Feriados: Si el lunes es feriado, el cierre se corre al martes
        dia_cierre_oficial = lunes_target
        while dia_cierre_oficial in feriados_cl:
            dia_cierre_oficial += timedelta(days=1)
            
        nombre_dia_cierre = "Lunes" if dia_cierre_oficial == lunes_target else "Martes (extendido por feriado)"
        
        # El candado evalúa la hora local y respeta la extensión si hubo festivo
        if not (viernes_prev <= hoy_dt <= dia_cierre_oficial):
            es_adicional = True
            st.warning(f"🔒 **Ventana de Planificación Cerrada:** El plazo oficial (Viernes a {nombre_dia_cierre}) ha expirado (Hora local: {ahora_chile.strftime('%H:%M:%S')}). Toda tarea ingresada ahora quedará etiquetada como **'Actividad Adicional'**.")

    tipo_actividad_actual = "Actividad Adicional" if es_adicional else "Programada"

    df_maestro = load_master_base()
    
    if df_maestro is not None:
        col_ajustador = 'Ajustador senior' if 'Ajustador senior' in df_maestro.columns else df_maestro.columns[9]
        ajustadores_validos = sorted(df_maestro[col_ajustador].dropna().unique())
        
        ajustador_seleccionado = st.selectbox("Identificación de Ajustador:", [""] + ajustadores_validos)
        
        if ajustador_seleccionado:
            casos_vigentes = df_maestro[(df_maestro[col_ajustador] == ajustador_seleccionado) & (df_maestro['Estado'] != 'Cerrado')].copy()
            
            if modo == "Mensual":
                casos_vigentes = casos_vigentes[casos_vigentes.apply(lambda x: calcular_tramo_mcl(x)[1], axis=1)]
                st.warning(f"📊 Filtro MCL Activo ({mes_opcion}): Mostrando exclusivamente siniestros complejos.")
            else:
                casos_vigentes = casos_vigentes[~casos_vigentes.apply(lambda x: calcular_tramo_mcl(x)[1], axis=1)]
                
            estados_maestros = sorted([str(x) for x in df_maestro['Estado'].dropna().unique() if str(x).strip()]) if 'Estado' in df_maestro.columns else ["Ajuste", "IFL", "Liquidación"]
            subestados_maestros = sorted([str(x) for x in df_maestro['Sub estado'].dropna().unique() if str(x).strip()]) if 'Sub estado' in df_maestro.columns else ["En Proceso", "Informe Preliminar", "Revisión Jefatura"]

            if modo == "Mensual":
                plan_historico, path_boveda = load_plan_mensual(ajustador_seleccionado, offset_months=offset_months)
            else:
                plan_historico, path_boveda = load_plan_semanal(ajustador_seleccionado, offset_weeks=offset_weeks)
            
            # --- VISIBILIDAD DEL PLAN VIGENTE ---
            if plan_historico:
                with st.expander(f"📋 Ver Plan Vigente ({len(plan_historico)} acciones registradas)", expanded=True):
                    df_ph = pd.DataFrame(plan_historico)
                    if not df_ph.empty:
                        df_show = df_ph[['tipo_actividad', 'categoria', 'numero_caso', 'accion', 'fecha_compromiso', 'estado_cumplimiento']].copy()
                        df_show = df_show.rename(columns={'tipo_actividad': 'Tipo', 'categoria': 'Categoría', 'numero_caso': 'Caso', 'accion': 'Acción', 'fecha_compromiso': 'Fecha', 'estado_cumplimiento': 'Estado'})
                        st.dataframe(df_show, use_container_width=True, hide_index=True)
                        
                with st.expander("🚨 Zona de Control: Modificar / Resetear Período Activo", expanded=False):
                    st.markdown("Si cometiste un error crítico, puedes vaciar el plan actual completo para volver a formularlo.")
                    if st.button("🗑️ ANULAR PLAN ACTUAL Y EMPEZAR DE CERO", key="btn_pánico_reset"):
                        try:
                            save_plan_actualizado(path_boveda, [])
                            st.success("¡Planificación anulada exitosamente! La pizarra está limpia.")
                            st.rerun()
                        except Exception as reset_err:
                            st.error(f"Error al ejecutar el reseteo: {reset_err}")

            plan_transaccional = []
            
            # --- HERENCIA MCL ---
            if modo == "Semanal":
                target_date = ahora_chile + timedelta(weeks=offset_weeks)
                target_month_id = target_date.strftime("%Y_%m")
                mcl_data, mcl_path = load_plan_mensual(ajustador_seleccionado, explicit_month_id=target_month_id)
                target_week_id = get_week_identifier(offset_weeks)
                
                mcl_pendientes = []
                for t in mcl_data:
                    try:
                        fec_obj = datetime.strptime(t['fecha_compromiso'], "%Y-%m-%d")
                        if fec_obj.strftime("%Y_W%W") == target_week_id and not t.get("agendado_semana"):
                            mcl_pendientes.append(t)
                    except: pass
                
                if mcl_pendientes:
                    st.markdown("---")
                    st.markdown('<div class="marco-gestion" style="border-left: 5px solid #d9534f;"><h4>🚨 Hitos MCL Heredados (Obligatorio asignar rango)</h4></div>', unsafe_allow_html=True)
                    st.info(f"Estos compromisos provienen de tu Planificador Mensual. Selecciona los días de ejecución (se omitirán fines de semana y feriados).")
                    
                    for idx, mcl_task in enumerate(mcl_pendientes):
                        c1, c2 = st.columns([3, 1])
                        with c1:
                            st.write(f"**Caso:** [{mcl_task['numero_caso']}] {mcl_task['asegurado']}")
                            st.write(f"**Entregable:** {mcl_task['accion']}")
                        with c2:
                            fec_obj = datetime.strptime(mcl_task['fecha_compromiso'], "%Y-%m-%d")
                            nueva_fecha_mcl = st.date_input(f"Días de ejecución:", value=(fec_obj, fec_obj), key=f"mcl_fec_{idx}")
                            
                        act_dates = []
                        if isinstance(nueva_fecha_mcl, (tuple, list)) and len(nueva_fecha_mcl) == 2:
                            s_d, e_d = nueva_fecha_mcl
                            delta = (e_d - s_d).days
                            for d in range(delta + 1):
                                dt = s_d + timedelta(days=d)
                                # Filtro de días hábiles y feriados
                                if dt.weekday() < 5 and dt not in feriados_cl:
                                    act_dates.append(dt)
                        elif isinstance(nueva_fecha_mcl, (tuple, list)) and len(nueva_fecha_mcl) == 1:
                            act_dates = [nueva_fecha_mcl[0]]
                        else:
                            act_dates = [nueva_fecha_mcl]

                        if not act_dates: 
                            act_dates = [s_d if 's_d' in locals() else nueva_fecha_mcl] # Paracaídas si eligen solo feriados

                        for dt in act_dates:
                            task_to_add = mcl_task.copy()
                            task_to_add['fecha_compromiso'] = dt.strftime("%Y-%m-%d")
                            task_to_add['id_mcl_origen'] = mcl_task['id_transaccion']
                            task_to_add['tipo_actividad'] = tipo_actividad_actual
                            plan_transaccional.append(task_to_add)

            st.markdown("---")
            st.header("1. Selección de Casos Operativos Regulares")
            st.info(f"Inventario Vigente: {len(casos_vigentes)} casos disponibles bajo este filtro.")
            
            def formato_caso_nickname(x):
                fila = casos_vigentes.loc[x]
                base_str = f"Caso {fila['Número de caso']} - {fila['Asegurado']}"
                if 'Nickname' in fila and pd.notna(fila['Nickname']) and str(fila['Nickname']).strip() != "":
                    return f"{base_str} - [{fila['Nickname']}]"
                return base_str

            selected_indices = st.multiselect(
                "Seleccione los casos que intervendrá:",
                options=casos_vigentes.index.tolist(),
                format_func=formato_caso_nickname
            )
            
            if selected_indices:
                st.markdown("---")
                st.header("2. Detalle de Acciones Operativas y Proyección de Estado")
                
                for idx in selected_indices:
                    fila = casos_vigentes.loc[idx]
                    caso_num = fila['Número de caso']
                    asegurado = fila['Asegurado']
                    estado_actual = str(fila['Estado']) if 'Estado' in fila and pd.notna(fila['Estado']) else "N/D"
                    subestado_actual = str(fila['Sub estado']) if 'Sub estado' in fila and pd.notna(fila['Sub estado']) else "N/D"
                    tramo, is_mcl = calcular_tramo_mcl(fila)
                    
                    honorarios_estimados = 0.0
                    try:
                        if len(casos_vigentes.columns) >= 67:
                            valor_bo = fila.iloc[66] 
                            honorarios_estimados = limpiar_monto_mcl(valor_bo)
                    except Exception: pass
                    
                    nickname = f" <span style='color:#004a99;'>[{fila['Nickname']}]</span>" if 'Nickname' in fila and pd.notna(fila['Nickname']) and str(fila['Nickname']).strip() != "" else ""
                    
                    with st.container():
                        st.markdown(f"""
                        <div class="marco-caso">
                            <h4>[{caso_num}] {asegurado}{nickname}</h4>
                            <p style="color:gray; font-size: 0.9em; margin-bottom: 5px;">
                                <b>Clasificación:</b> {tramo} | <b>Estado Actual:</b> {estado_actual} | <b>Sub-estado Actual:</b> {subestado_actual}
                            </p>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        col_est, col_sub = st.columns(2)
                        with col_est:
                            opts_est = sorted(list(set(estados_maestros + ([estado_actual] if estado_actual != "N/D" else []))))
                            default_est_idx = opts_est.index(estado_actual) if estado_actual in opts_est else 0
                            estado_proyectado = st.selectbox(f"Proyectar Estado Final:", opts_est, index=default_est_idx, key=f"est_proj_{idx}")
                            
                        with col_sub:
                            opts_sub = sorted(list(set(subestados_maestros + ([subestado_actual] if subestado_actual != "N/D" else []))))
                            default_sub_idx = opts_sub.index(subestado_actual) if subestado_actual in opts_sub else 0
                            subestado_proyectado = st.selectbox(f"Proyectar Sub-estado Final:", opts_sub, index=default_sub_idx, key=f"sub_proj_{idx}")

                        num_actividades = st.number_input(f"Cantidad de actividades para el caso {caso_num}:", min_value=1, max_value=15, value=3, key=f"num_act_{idx}")
                        
                        for i in range(1, int(num_actividades) + 1):
                            colA, colB, colC = st.columns([2, 2, 1])
                            with colA:
                                cat_accion = st.selectbox(f"Categoría Acción {i}:", [""] + list(CATALOGO_ACCIONES.keys()), key=f"cat_{idx}_{i}")
                            with colB:
                                accion_final = ""
                                if cat_accion:
                                    if cat_accion == "Otra Acción (Manual)":
                                        accion_final = st.text_input(f"Describa la acción {i}:", key=f"man_{idx}_{i}")
                                    else:
                                        sub_accion = st.selectbox(f"Detalle Acción {i}:", [""] + CATALOGO_ACCIONES[cat_accion], key=f"sub_{idx}_{i}")
                                        if sub_accion == "Otro / Manual":
                                            texto_manual = st.text_input(f"Especifique el detalle {i}:", key=f"man_{idx}_{i}")
                                            if texto_manual: accion_final = f"{cat_accion} - {texto_manual}"
                                        elif sub_accion:
                                            accion_final = f"{cat_accion} - {sub_accion}"
                            with colC:
                                fecha_compromiso_range = st.date_input(f"Rango ejecución {i}:", value=(ahora_chile.date(), ahora_chile.date()), key=f"fecha_{idx}_{i}")
                            
                            if accion_final.strip():
                                act_dates = []
                                if isinstance(fecha_compromiso_range, (tuple, list)) and len(fecha_compromiso_range) == 2:
                                    s_d, e_d = fecha_compromiso_range
                                    delta = (e_d - s_d).days
                                    for d in range(delta + 1):
                                        dt = s_d + timedelta(days=d)
                                        # Colador de Fines de semana y Feriados
                                        if dt.weekday() < 5 and dt not in feriados_cl:
                                            act_dates.append(dt)
                                elif isinstance(fecha_compromiso_range, (tuple, list)) and len(fecha_compromiso_range) == 1:
                                    act_dates = [fecha_compromiso_range[0]]
                                else:
                                    act_dates = [fecha_compromiso_range]

                                if not act_dates: act_dates = [s_d if 's_d' in locals() else fecha_compromiso_range]

                                for dt in act_dates:
                                    plan_transaccional.append({
                                        "id_transaccion": str(uuid.uuid4()),
                                        "tipo_plan": modo,
                                        "tipo_actividad": tipo_actividad_actual,
                                        "categoria": "Operativa",
                                        "numero_caso": str(caso_num),
                                        "asegurado": str(asegurado),
                                        "tramo_uf": tramo,
                                        "honorarios_estimados": honorarios_estimados,
                                        "estado_proyectado": estado_proyectado,
                                        "subestado_proyectado": subestado_proyectado,
                                        "accion": accion_final,
                                        "fecha_compromiso": dt.strftime("%Y-%m-%d"),
                                        "estado_cumplimiento": "Pendiente",
                                        "fecha_planificacion": ahora_chile.strftime("%Y-%m-%d %H:%M:%S")
                                    })
                                
            st.markdown("---")
            st.header("3. Acciones de Gestión")
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown('<div class="marco-gestion"><h4>🤝 Gestión Comercial</h4></div>', unsafe_allow_html=True)
                num_comercial = st.number_input("Cantidad de gestiones comerciales:", min_value=0, max_value=15, value=1, key="num_comercial")
                for i in range(1, int(num_comercial) + 1):
                    c_acc, c_fec = st.columns([2, 1])
                    with c_acc:
                        acc_com = st.text_input(f"Detalle gestión {i}:", placeholder="Reuniones, visitas a corredoras...", key=f"txt_com_{i}")
                    with c_fec:
                        fec_com_range = st.date_input(f"Rango {i}:", value=(ahora_chile.date(), ahora_chile.date()), key=f"fec_com_{i}")
                    
                    if acc_com.strip():
                        com_dates = []
                        if isinstance(fec_com_range, (tuple, list)) and len(fec_com_range) == 2:
                            s_d, e_d = fec_com_range
                            delta = (e_d - s_d).days
                            for d in range(delta + 1):
                                dt = s_d + timedelta(days=d)
                                if dt.weekday() < 5 and dt not in feriados_cl:
                                    com_dates.append(dt)
                        elif isinstance(fec_com_range, (tuple, list)) and len(fec_com_range) == 1:
                            com_dates = [fec_com_range[0]]
                        else:
                            com_dates = [fec_com_range]

                        if not com_dates: com_dates = [s_d if 's_d' in locals() else fec_com_range]

                        for dt in com_dates:
                            plan_transaccional.append({
                                "id_transaccion": str(uuid.uuid4()), "tipo_plan": modo, "tipo_actividad": tipo_actividad_actual, 
                                "categoria": "Gestión Comercial", "numero_caso": "N/A", "asegurado": "N/A", "tramo_uf": "N/A", 
                                "honorarios_estimados": 0.0, "estado_proyectado": "N/A", "subestado_proyectado": "N/A", 
                                "accion": acc_com.strip(), "fecha_compromiso": dt.strftime("%Y-%m-%d"), 
                                "estado_cumplimiento": "Pendiente", "fecha_planificacion": ahora_chile.strftime("%Y-%m-%d %H:%M:%S")
                            })
            
            with col2:
                st.markdown('<div class="marco-gestion"><h4>⚙️ Gestión Administrativa</h4></div>', unsafe_allow_html=True)
                num_admin = st.number_input("Cantidad de gestiones administrativas:", min_value=0, max_value=15, value=1, key="num_admin")
                for i in range(1, int(num_admin) + 1):
                    c_acc, c_fec = st.columns([2, 1])
                    with c_acc:
                        acc_adm = st.text_input(f"Detalle gestión {i}:", placeholder="Capacitaciones, comités...", key=f"txt_adm_{i}")
                    with c_fec:
                        fec_adm_range = st.date_input(f"Rango {i}:", value=(ahora_chile.date(), ahora_chile.date()), key=f"fec_adm_{i}")
                    
                    if acc_adm.strip():
                        adm_dates = []
                        if isinstance(fec_adm_range, (tuple, list)) and len(fec_adm_range) == 2:
                            s_d, e_d = fec_adm_range
                            delta = (e_d - s_d).days
                            for d in range(delta + 1):
                                dt = s_d + timedelta(days=d)
                                if dt.weekday() < 5 and dt not in feriados_cl:
                                    adm_dates.append(dt)
                        elif isinstance(fec_adm_range, (tuple, list)) and len(fec_adm_range) == 1:
                            adm_dates = [fec_adm_range[0]]
                        else:
                            adm_dates = [fec_adm_range]

                        if not adm_dates: adm_dates = [s_d if 's_d' in locals() else fec_adm_range]

                        for dt in adm_dates:
                            plan_transaccional.append({
                                "id_transaccion": str(uuid.uuid4()), "tipo_plan": modo, "tipo_actividad": tipo_actividad_actual, 
                                "categoria": "Gestión Administrativa", "numero_caso": "N/A", "asegurado": "N/A", "tramo_uf": "N/A", 
                                "honorarios_estimados": 0.0, "estado_proyectado": "N/A", "subestado_proyectado": "N/A", 
                                "accion": acc_adm.strip(), "fecha_compromiso": dt.strftime("%Y-%m-%d"), 
                                "estado_cumplimiento": "Pendiente", "fecha_planificacion": ahora_chile.strftime("%Y-%m-%d %H:%M:%S")
                            })

            st.markdown("---")
            if len(plan_transaccional) > 0:
                st.info(f"Se consolidaron **{len(plan_transaccional)} transacciones** para guardar.")
                if st.button(f"💾 GUARDAR REGISTROS ({tipo_actividad_actual.upper()})"):
                    try:
                        if modo == "Mensual":
                            plan_existente, filepath = load_plan_mensual(ajustador_seleccionado, offset_months=offset_months)
                            save_plan_actualizado(filepath, plan_existente + plan_transaccional)
                            st.success(f"Plan Mensual MCL ({mes_opcion}) actualizado exitosamente.")
                        else:
                            plan_existente, filepath = load_plan_semanal(ajustador_seleccionado, offset_weeks=offset_weeks)
                            
                            if 'mcl_data' in locals() and mcl_data:
                                mcl_ids_agendados = [t['id_mcl_origen'] for t in plan_transaccional if 'id_mcl_origen' in t]
                                for t in mcl_data:
                                    if t['id_transaccion'] in mcl_ids_agendados:
                                        t['agendado_semana'] = True
                                save_plan_actualizado(mcl_path, mcl_data) 
                                
                            save_plan_actualizado(filepath, plan_existente + plan_transaccional)
                            st.success(f"Registro exitoso para la {semana_opcion}. Documento respaldado.")
                            st.rerun()
                    except Exception as e:
                        st.error(f"Error al guardar: {e}")
            elif selected_indices or int(num_comercial) > 0 or int(num_admin) > 0:
                st.warning("Complete el detalle de las acciones o seleccione casos válidos para guardar.")
    else:
        st.info("Módulo en espera: Suba el archivo 'Reporte de acciones' en el panel izquierdo.")

# ---------------------------------------------------------
# BLOQUE 3: VISTA - PROGRAMA DIARIO (EJECUCIÓN Y NO PROGRAMADOS)
# VERSIÓN: 2.4.3 (Autocompletado con Herencia Financiera UF)
# ---------------------------------------------------------
def vista_diario():
    import json
    import os
    import uuid
    import pandas as pd
    import streamlit as st
    from datetime import datetime
    
    st.title("☀️ Ejecución y Cumplimiento")
    hoy_str = datetime.now().strftime("%Y-%m-%d")
    st.markdown(f"**Fecha actual:** {datetime.now().strftime('%A, %d de %B de %Y')}")
    
    week_id = get_week_identifier()
    
    # --- MOTOR DE SINCRONIZACIÓN AUTOMÁTICA (Rescate de Nube) ---
    archivos_locales = [f for f in os.listdir(PERSISTENCE_DIR) if f.startswith('plan_') and f.endswith('.json')]
    if not archivos_locales:
        with st.spinner("Sincronizando planes desde la nube corporativa..."):
            sheet = get_google_sheet()
            if sheet:
                try:
                    registros = sheet.get_all_records()
                    for fila in registros:
                        fname = str(fila.get('Archivo', ''))
                        if fname.endswith('.json') and fname != "BASE_MAESTRA.json":
                            fpath = os.path.join(PERSISTENCE_DIR, fname)
                            try:
                                datos_json = json.loads(fila.get('JSON_Data', '[]'))
                                with open(fpath, 'w', encoding='utf-8') as f:
                                    json.dump(datos_json, f, ensure_ascii=False, indent=4)
                            except Exception:
                                pass
                except Exception:
                    pass

    # --- MOTOR DE BÚSQUEDA GLOBAL DE AJUSTADORES (Base Maestra) ---
    df_maestro = load_master_base()
    ajustadores_validos = []
    
    if df_maestro is not None:
        col_ajustador = 'Ajustador senior' if 'Ajustador senior' in df_maestro.columns else df_maestro.columns[9]
        ajustadores_validos = sorted(df_maestro[col_ajustador].dropna().unique())
        
    if not ajustadores_validos:
        st.info("⚠️ No se pudo cargar la lista de ajustadores desde la Base Maestra.")
        return

    ajustador_input = st.selectbox("Seleccione su nombre de Ajustador:", [""] + ajustadores_validos)
    
    if ajustador_input:
        plan_data, filepath = load_plan_semanal(ajustador_input)
        
        # --- CARGA DE CASOS PARA AUTOCOMPLETADO (INCLUYE UF) ---
        opciones_casos = ["Gestión Manual (Caso fuera de sistema)"]
        dict_casos = {}
        if df_maestro is not None and not df_maestro.empty:
            casos_ajustador = df_maestro[(df_maestro[col_ajustador] == ajustador_input) & (df_maestro['Estado'] != 'Cerrado')]
            for _, row in casos_ajustador.iterrows():
                num = str(row.get('Número de caso', '')).strip()
                if num:
                    aseg = str(row.get('Asegurado', '')).strip()
                    nick = str(row.get('Nickname', '')).strip()
                    
                    # Rescate de honorarios para herencia automática
                    tramo, is_mcl = calcular_tramo_mcl(row)
                    honorarios = 0.0
                    try:
                        if len(df_maestro.columns) >= 67:
                            valor_bo = row.iloc[66]
                            honorarios = limpiar_monto_mcl(valor_bo)
                    except: pass
                    
                    lbl = f"[{num}] {aseg}"
                    if nick:
                        lbl += f" - [{nick}]"
                    opciones_casos.append(lbl)
                    dict_casos[lbl] = {
                        "caso": num, 
                        "asegurado": aseg,
                        "tramo": tramo,
                        "honorarios": honorarios
                    }

        # --- MÓDULO DE ACTIVIDADES NO PROGRAMADAS ---
        st.markdown("---")
        with st.expander("➕ REGISTRAR ACTIVIDADES NO PROGRAMADAS (Urgencias / Fuera de Plan)", expanded=False):
            st.info("Utilice este módulo para reportar gestiones inmediatas que no estaban en su planificación original. Estas impactan positivamente en su métrica de cumplimiento global.")
            
            num_np = st.number_input("Cantidad de urgencias a registrar ahora:", min_value=1, max_value=15, value=1, key="num_np_diario")
            
            nuevas_actividades = []
            
            for i in range(1, int(num_np) + 1):
                st.markdown(f"**Urgencia {i}**")
                
                seleccion_caso = st.selectbox(f"Asociar a un caso vigente {i}:", options=opciones_casos, key=f"sel_caso_{i}")
                
                colNP1, colNP2 = st.columns(2)
                with colNP1:
                    if seleccion_caso == "Gestión Manual (Caso fuera de sistema)":
                        np_caso = st.text_input(f"Número de Caso (o Ref) {i}:", key=f"np_caso_{i}")
                        np_asegurado = st.text_input(f"Asegurado {i}:", key=f"np_aseg_{i}")
                        np_tramo = "N/D"
                        np_honorarios = 0.0
                    else:
                        np_caso = dict_casos[seleccion_caso]["caso"]
                        np_asegurado = dict_casos[seleccion_caso]["asegurado"]
                        np_tramo = dict_casos[seleccion_caso]["tramo"]
                        np_honorarios = dict_casos[seleccion_caso]["honorarios"]
                        
                        st.text_input(f"Número de Caso {i}:", value=np_caso, disabled=True, key=f"np_caso_dis_{i}")
                        st.text_input(f"Asegurado {i}:", value=np_asegurado, disabled=True, key=f"np_aseg_dis_{i}")
                        
                with colNP2:
                    np_accion = st.text_input(f"Acción Ejecutada {i}:", key=f"np_acc_{i}")
                    np_fecha = st.date_input(f"Fecha de Ejecución {i}:", value=datetime.now(), key=f"np_fec_{i}")
                
                if np_caso.strip() and np_accion.strip():
                    nuevas_actividades.append({
                        "id_transaccion": str(uuid.uuid4()),
                        "tipo_plan": "Diario",
                        "tipo_actividad": "Actividad Adicional",
                        "categoria": "Operativa",
                        "numero_caso": str(np_caso),
                        "asegurado": str(np_asegurado),
                        "tramo_uf": np_tramo,
                        "honorarios_estimados": float(np_honorarios), 
                        "estado_proyectado": "N/D",
                        "subestado_proyectado": "N/D",
                        "accion": np_accion,
                        "fecha_compromiso": np_fecha.strftime("%Y-%m-%d"),
                        "estado_cumplimiento": "Realizado", 
                        "fecha_ejecucion": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        "fecha_planificacion": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    })
                st.markdown("<hr style='margin-top: 5px; margin-bottom: 15px;'>", unsafe_allow_html=True)
            
            if st.button("💾 Guardar Actividades No Programadas"):
                if len(nuevas_actividades) > 0:
                    plan_data.extend(nuevas_actividades)
                    save_plan_actualizado(filepath, plan_data)
                    st.success(f"¡{len(nuevas_actividades)} actividades no programadas incorporadas al reporte general exitosamente!")
                    st.rerun()
                else:
                    st.warning("⚠️ Debe ingresar al menos el Número de Caso y la Acción Ejecutada en alguna de las filas para poder guardar.")

        if not plan_data:
            st.warning(f"⚠️ No se encontraron compromisos agendados para **{ajustador_input}** en la semana en curso. Sin embargo, puede registrar sus urgencias en el módulo superior.")
        else:
            st.success("Plan Operativo sincronizado correctamente.")
            st.markdown("---")
            
            # --- MOTOR DE CÁLCULO DE HONORARIOS Y TAREAS ---
            tareas_hoy = []
            tareas_resto = []
            uf_proyectadas_hoy = 0.0
            uf_ejecutadas_hoy = 0.0
            uf_proyectadas_semana = 0.0
            uf_ejecutadas_semana = 0.0
            
            for idx, tarea in enumerate(plan_data):
                tarea_con_indice = tarea.copy()
                tarea_con_indice['_posicion_original'] = idx
                
                try:
                    uf_tarea = float(tarea.get("honorarios_estimados", 0.0))
                except:
                    uf_tarea = 0.0
                    
                es_realizado = (tarea.get("estado_cumplimiento") == "Realizado")
                
                uf_proyectadas_semana += uf_tarea
                if es_realizado:
                    uf_ejecutadas_semana += uf_tarea
                
                if tarea.get("fecha_compromiso") == hoy_str:
                    tareas_hoy.append(tarea_con_indice)
                    uf_proyectadas_hoy += uf_tarea
                    if es_realizado:
                        uf_ejecutadas_hoy += uf_tarea
                else:
                    tareas_resto.append(tarea_con_indice)
            
            total_tareas = len(plan_data)
            tareas_completadas = sum(1 for t in plan_data if t.get("estado_cumplimiento") == "Realizado")
            cambios_realizados = False
            
            # --- TABLERO VISUAL DE HONORARIOS ---
            st.subheader("📊 Rendimiento Financiero del Plan")
            col_met1, col_met2, col_met3, col_met4 = st.columns(4)
            col_met1.metric("UF Proyectadas Hoy", f"{uf_proyectadas_hoy:,.2f}")
            col_met2.metric("UF Ejecutadas Hoy", f"{uf_ejecutadas_hoy:,.2f}")
            col_met3.metric("UF Proyectadas Semana", f"{uf_proyectadas_semana:,.2f}")
            col_met4.metric("UF Ejecutadas Semana", f"{uf_ejecutadas_semana:,.2f}")
            st.markdown("---")
            
            with st.form(key="form_cumplimiento_estructurado"):
                if tareas_hoy:
                    st.subheader("🔥 Prioridad para Hoy (Compromisos del Día)")
                    for t in tareas_hoy:
                        pos = t['_posicion_original']
                        estado_actual = t.get("estado_cumplimiento", "Pendiente")
                        es_realizado = (estado_actual == "Realizado")
                        clase_css = "tarea-marco tarea-realizada" if es_realizado else "tarea-marco"
                        icono = "✅" if es_realizado else "⚡"
                        tipo_act = f" [{t.get('tipo_actividad', 'Programada').upper()}]"
                        uf_txt = f" | 💰 {float(t.get('honorarios_estimados', 0.0)):,.2f} UF"
                        
                        st.markdown(f'<div class="{clase_css}">', unsafe_allow_html=True)
                        if t["categoria"] == "Operativa":
                            est_p = t.get('estado_proyectado', 'N/D')
                            sub_p = t.get('subestado_proyectado', 'N/D')
                            st.markdown(f"**{icono} CASO [{t['numero_caso']}]** - {t['asegurado']} | *Tramo: {t['tramo_uf']}* | *Proyectado: {est_p} ({sub_p})*{tipo_act}{uf_txt}")
                        else:
                            st.markdown(f"**{icono} {t['categoria'].upper()}**{tipo_act}{uf_txt}")
                        st.markdown(f"**Entregable:** {t['accion']}")
                        
                        nuevo_estado = st.checkbox(f"Marcar como ejecutado", value=es_realizado, key=f"chk_hoy_{t['id_transaccion']}")
                        nuevo_texto_estado = "Realizado" if nuevo_estado else "Pendiente"
                        if nuevo_texto_estado != estado_actual:
                            plan_data[pos]["estado_cumplimiento"] = nuevo_texto_estado
                            plan_data[pos]["fecha_ejecucion"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S") if nuevo_estado else ""
                            cambios_realizados = True
                        st.markdown('</div>', unsafe_allow_html=True)
                else:
                    st.info("💡 No tienes actividades agendadas específicamente para la fecha de hoy. Abajo se despliega tu planificación extendida.")

                if tareas_resto:
                    st.subheader("📅 Resto de la Planificación (Semanal y Mensual)")
                    for t in tareas_resto:
                        pos = t['_posicion_original']
                        estado_actual = t.get("estado_cumplimiento", "Pendiente")
                        es_realizado = (estado_actual == "Realizado")
                        clase_css = "tarea-marco tarea-realizada" if es_realizado else "tarea-marco"
                        icono = "✅" if es_realizado else "⏳"
                        tipo_act = f" [{t.get('tipo_actividad', 'Programada').upper()}]"
                        uf_txt = f" | 💰 {float(t.get('honorarios_estimados', 0.0)):,.2f} UF"
                        
                        st.markdown(f'<div class="{clase_css}">', unsafe_allow_html=True)
                        if t["categoria"] == "Operativa":
                            est_p = t.get('estado_proyectado', 'N/D')
                            st.markdown(f"**{icono} CASO [{t['numero_caso']}]** - {t['asegurado']} | *Compromiso: {t['fecha_compromiso']}* | *Proyectado: {est_p}*{tipo_act}{uf_txt}")
                        else:
                            st.markdown(f"**{icono} {t['categoria'].upper()}** | *Compromiso: {t['fecha_compromiso']}*{tipo_act}{uf_txt}")
                        st.markdown(f"**Entregable:** {t['accion']}")
                        
                        nuevo_estado = st.checkbox(f"Marcar como ejecutado", value=es_realizado, key=f"chk_rest_{t['id_transaccion']}")
                        nuevo_texto_estado = "Realizado" if nuevo_estado else "Pendiente"
                        if nuevo_texto_estado != estado_actual:
                            plan_data[pos]["estado_cumplimiento"] = nuevo_texto_estado
                            plan_data[pos]["fecha_ejecucion"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S") if nuevo_estado else ""
                            cambios_realizados = True
                        st.markdown('</div>', unsafe_allow_html=True)
                
                st.markdown('<div class="btn-guardar">', unsafe_allow_html=True)
                submit_button = st.form_submit_button(label="💾 ACTUALIZAR CUMPLIMIENTO DIARIO")
                st.markdown('</div>', unsafe_allow_html=True)
            
            if submit_button:
                if cambios_realizados:
                    try:
                        save_plan_actualizado(filepath, plan_data)
                        st.success("¡Base de datos local actualizada y respaldada en la nube corporativa!")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error al escribir en el disco/nube: {e}")
                else:
                    st.info("No se detectaron cambios.")
            
            st.markdown("---")
            progreso = int((tareas_completadas / total_tareas) * 100) if total_tareas > 0 else 0
            st.progress(progreso)
            st.caption(f"Avance de cumplimiento global: {tareas_completadas} de {total_tareas} tareas realizadas ({progreso}%).")


# =========================================================
# MÓDULO 4: REPORTE DE JEFATURA (GANTT, DASHBOARD Y OPERACIONAL)
# VERSIÓN: 4.9.2 (Arquitectura Modular Segura y Blindaje de Variables)
# =========================================================

# ---------------------------------------------------------
# BLOQUE 4.1: UTILIDADES GRÁFICAS Y FORMATOS DE TABLA
# ---------------------------------------------------------
def obtener_color_kpi(valor):
    if valor <= 50: 
        return "#d9534f" # Rojo
    elif valor < 80: 
        return "#f0ad4e" # Amarillo
    else: 
        return "#28a745" # Verde

def crear_velocimetro(valor, titulo, inverso=False):
    import plotly.graph_objects as go
    if inverso:
        pasos = [
            {'range': [0, 20], 'color': "#28a745"},
            {'range': [20, 50], 'color': "#f0ad4e"},
            {'range': [50, 100], 'color': "#d9534f"}
        ]
    else:
        pasos = [
            {'range': [0, 50], 'color': "#d9534f"},
            {'range': [50, 79.99], 'color': "#f0ad4e"},
            {'range': [80, 100], 'color': "#28a745"}
        ]
        
    fig = go.Figure(go.Indicator(
        mode = "gauge+number",
        value = valor,
        title = {'text': titulo, 'font': {'size': 18}},
        number = {'suffix': "%", 'font': {'size': 26, 'color': "black", 'weight': 'bold'}},
        gauge = {
            'axis': {'range': [None, 100], 'tickwidth': 1, 'tickcolor': "black"},
            'bar': {'color': "black", 'thickness': 0.15},
            'bgcolor': "white",
            'borderwidth': 2,
            'bordercolor': "gray",
            'steps': pasos,
        }
    ))
    fig.update_layout(height=250, margin=dict(l=20, r=20, t=50, b=20))
    return fig

def formatear_cabecera_tabla(table, bg_color="003366"):
    from docx.shared import RGBColor
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    hdr_cells = table.rows[0].cells
    for cell in hdr_cells:
        if cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
        else:
            run = cell.paragraphs[0].add_run(cell.text)
        
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
        cell._tc.get_or_add_tcPr().append(shading_elm)

# ---------------------------------------------------------
# BLOQUE 4.2: MOTOR DE SINCRONIZACIÓN Y EXTRACCIÓN DE DATOS
# ---------------------------------------------------------
def sincronizar_y_cargar_datos(forzar_sync, dias_semana_target):
    import os
    import json
    import pandas as pd
    import streamlit as st

    archivos_existentes = [f for f in os.listdir(PERSISTENCE_DIR) if f.startswith('plan_') and f.endswith('.json')]
    
    if forzar_sync or len(archivos_existentes) == 0:
        with st.spinner("Descargando planes de todos los ajustadores desde Google Sheets..."):
            sheet = get_google_sheet()
            if sheet:
                try:
                    registros = sheet.get_all_records()
                    for fila in registros:
                        fname = str(fila.get('Archivo', ''))
                        if fname.endswith('.json') and fname != "BASE_MAESTRA.json":
                            fpath = os.path.join(PERSISTENCE_DIR, fname)
                            try:
                                datos_json = json.loads(fila.get('JSON_Data', '[]'))
                                with open(fpath, 'w', encoding='utf-8') as f:
                                    json.dump(datos_json, f, ensure_ascii=False, indent=4)
                            except Exception:
                                pass
                    if forzar_sync:
                        st.success("¡Sincronización completada! Todos los planes están actualizados.")
                except Exception as e:
                    st.warning(f"Error al sincronizar con la nube: {e}")
                    
    archivos_json = [f for f in os.listdir(PERSISTENCE_DIR) if f.endswith('.json') and f != "BASE_MAESTRA.json"]
    
    df_maestro = load_master_base()
    diccionario_nicknames = {}
    ajustadores_validos = []
    
    if df_maestro is not None:
        if 'Número de caso' in df_maestro.columns and 'Nickname' in df_maestro.columns:
            for _, row in df_maestro.iterrows():
                if pd.notna(row['Número de caso']) and str(row['Número de caso']).strip() != "":
                    if pd.notna(row['Nickname']):
                        diccionario_nicknames[str(row['Número de caso'])] = str(row['Nickname'])
                    else:
                        diccionario_nicknames[str(row['Número de caso'])] = ""
        
        if 'Ajustador senior' in df_maestro.columns:
            col_ajustador = 'Ajustador senior'
        else:
            col_ajustador = df_maestro.columns[9]
            
        ajustadores_validos = sorted(df_maestro[col_ajustador].dropna().unique())

    datos_consolidados = []
    for archivo in archivos_json:
        partes_nombre = archivo.replace('plan_mensual_mcl_', '').replace('plan_', '').replace('.json', '').split('_20')
        nombre_ajustador = partes_nombre[0].replace('_', ' ')
        
        filepath = os.path.join(PERSISTENCE_DIR, archivo)
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                plan = json.load(f)
                for tarea in plan:
                    tarea['Ajustador'] = nombre_ajustador
                    numero_caso_str = str(tarea.get('numero_caso', ''))
                    tarea['Nick Name'] = diccionario_nicknames.get(numero_caso_str, '')
                    datos_consolidados.append(tarea)
        except Exception:
            pass

    df_week = pd.DataFrame()
    df_raw = pd.DataFrame()
    
    if datos_consolidados:
        df_raw = pd.DataFrame(datos_consolidados)
        df_raw['fecha_filtro'] = pd.to_datetime(df_raw['fecha_compromiso'], errors='coerce').dt.date
        df_week = df_raw[df_raw['fecha_filtro'].isin(dias_semana_target)].copy()
        
        if not df_week.empty:
            if 'honorarios_estimados' not in df_week.columns:
                df_week['honorarios_estimados'] = 0.0
            df_week['honorarios_estimados'] = pd.to_numeric(df_week['honorarios_estimados'], errors='coerce').fillna(0.0)
            
            if 'tramo_uf' not in df_week.columns:
                df_week['tramo_uf'] = 'N/D'
                
    return df_week, df_raw, ajustadores_validos

# ---------------------------------------------------------
# BLOQUE 4.3: VISTA - DASHBOARD EJECUTIVO (BI)
# VERSIÓN: 4.3.10 (Paginación Inteligente y Optimización de Fuente)
# ---------------------------------------------------------
def renderizar_dashboard_ejecutivo(df_week, target_week_id, week_id_obj):
    import io
    import pandas as pd
    import streamlit as st
    import plotly.graph_objects as go
    import plotly.express as px
    from openpyxl import Workbook
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    
    if df_week.empty:
        st.info("No hay datos cargados para generar el Dashboard en esta semana.")
        return

    df_realizados = df_week[df_week['estado_cumplimiento'] == 'Realizado'].copy()
    
    # --- 1. RESUMEN EJECUTIVO FINANCIERO ---
    st.markdown('<div class="marco-gestion" style="border-left: 5px solid #003366;"><h4>💰 Cuadrante 1: Valorización de Cartera y Facturación Proyectada</h4></div>', unsafe_allow_html=True)
    cond_facturable = df_realizados['accion'].str.contains('Informe Final de Liquidación|Carta de Cobertura \(Rechazo\)', case=False, na=False)
    
    uf_facturables_caja = df_realizados[cond_facturable]['honorarios_estimados'].sum()
    
    cond_wip = (df_realizados['categoria'] == 'Operativa') & (~cond_facturable)
    uf_traccionadas_wip = df_realizados[cond_wip]['honorarios_estimados'].sum()
    
    c_fin1, c_fin2 = st.columns(2)
    with c_fin1:
        st.metric("Ingreso Efectivo Facturable (Cierres/Rechazos)", f"{uf_facturables_caja:,.2f} UF", delta="Directo a Caja", delta_color="normal")
    with c_fin2:
        st.metric("Valor Potencial Traccionado (WIP / Proceso)", f"{uf_traccionadas_wip:,.2f} UF", delta="Esfuerzo Operativo", delta_color="off")
    st.markdown("---")
    
    # --- 2. KPIS DE CUMPLIMIENTO CON VELOCÍMETROS PLOTLY (PARA WEB) ---
    st.markdown('<div class="marco-gestion" style="border-left: 5px solid #17a2b8;"><h4>⏱️ Cuadrante 2: Métricas de Cumplimiento y Carga Operativa</h4></div>', unsafe_allow_html=True)
    
    t_planificadas = len(df_week[df_week['tipo_actividad'] == 'Programada'])
    t_planificadas_hechas = len(df_realizados[df_realizados['tipo_actividad'] == 'Programada'])
    t_urgencias = len(df_realizados[df_realizados['tipo_actividad'] == 'Actividad Adicional'])
    
    if t_planificadas > 0:
        adherencia = (t_planificadas_hechas / t_planificadas) * 100
    else:
        adherencia = 0
        
    total_esfuerzo = t_planificadas_hechas + t_urgencias
    
    if total_esfuerzo > 0:
        ratio_plan = (t_planificadas_hechas / total_esfuerzo) * 100
        ratio_urg = (t_urgencias / total_esfuerzo) * 100
    else:
        ratio_plan = 0
        ratio_urg = 0
    
    fig_adh_dash = crear_velocimetro(adherencia, "Adherencia al Plan")
    fig_pro_dash = crear_velocimetro(ratio_plan, "Ratio Planificado")
    
    c_kpi1, c_kpi2 = st.columns(2)
    with c_kpi1: 
        st.plotly_chart(fig_adh_dash, use_container_width=True, key=f"dash_adh_{target_week_id}")
    with c_kpi2: 
        st.plotly_chart(fig_pro_dash, use_container_width=True, key=f"dash_pro_{target_week_id}")
    st.markdown("---")
    
    # --- 3. RADAR DE CASOS MCL ---
    st.markdown('<div class="marco-gestion" style="border-left: 5px solid #d9534f;"><h4>🏆 Radar de Casos MCL (Major and Complex Losses)</h4></div>', unsafe_allow_html=True)
    cond_mcl = df_realizados['tramo_uf'].str.contains('MCL|> 5', case=False, na=False)
    df_mcl = df_realizados[cond_mcl][['Ajustador', 'numero_caso', 'asegurado', 'accion', 'honorarios_estimados']].copy()
    
    df_mcl_mostrar = pd.DataFrame(columns=['Caso', 'Asegurado', 'Gestión MCL', 'Hon UF'])
    
    if not df_mcl.empty:
        df_mcl = df_mcl.groupby(['Ajustador', 'numero_caso', 'asegurado', 'accion'], as_index=False).agg({'honorarios_estimados': 'max'})
        
        df_mcl['accion'] = df_mcl['accion'].str.replace('Preparar Informe - ', '', regex=False)
        df_mcl['accion'] = df_mcl['accion'].str.replace('Reunión - ', '', regex=False)
        df_mcl['honorarios_estimados'] = df_mcl['honorarios_estimados'].apply(lambda x: f"{x:,.2f}")
        
        df_mcl_mostrar = df_mcl.rename(columns={
            'numero_caso': 'Caso', 
            'asegurado': 'Asegurado', 
            'accion': 'Gestión MCL', 
            'honorarios_estimados': 'Hon UF'
        })
        st.dataframe(df_mcl_mostrar, use_container_width=True, hide_index=True)
    else:
        st.info("Sin movimientos reportados en casos de tramo MCL esta semana.")
    st.markdown("---")

    # --- 4. INVENTARIO DE PRODUCCIÓN ---
    st.markdown('<div class="marco-gestion" style="border-left: 5px solid #28a745;"><h4>🏭 Cuadrante 3: Inventario de Producción (Entregables de Valor)</h4></div>', unsafe_allow_html=True)
    cond_entregables = df_realizados['accion'].str.contains('Preparar Informe|Presentación pptx|Presentacion on line', case=False, na=False)
    df_entregables = df_realizados[cond_entregables][['Ajustador', 'numero_caso', 'Nick Name', 'asegurado', 'accion', 'honorarios_estimados']].copy()
    
    df_mostrar = pd.DataFrame(columns=['Ajustador', 'Caso', 'Nick Name', 'Asegurado', 'Tipo de Entregable', 'Hon UF'])
    
    if not df_entregables.empty:
        df_entregables = df_entregables.groupby(['Ajustador', 'numero_caso', 'Nick Name', 'asegurado', 'accion'], as_index=False).agg({'honorarios_estimados': 'max'})
        
        df_entregables['accion'] = df_entregables['accion'].str.replace('Preparar Informe - ', '', regex=False)
        df_entregables['accion'] = df_entregables['accion'].str.replace('Reunión - ', '', regex=False)
        
        # --- GRÁFICOS DE TORTA (UI Web y Orden de Magnitud) ---
        df_pie_qty = df_entregables['accion'].value_counts().reset_index()
        df_pie_qty.columns = ['Entregable', 'Cantidad']
        df_pie_qty = df_pie_qty.sort_values(by='Cantidad', ascending=False)
        
        df_pie_uf = df_entregables.groupby('accion')['honorarios_estimados'].sum().reset_index()
        df_pie_uf.columns = ['Entregable', 'UF']
        df_pie_uf = df_pie_uf.sort_values(by='UF', ascending=False)
        
        fig_qty = px.pie(df_pie_qty, values='Cantidad', names='Entregable', title='Distribución por Cantidad de Entregables', hole=0.3)
        fig_qty.update_traces(textposition='inside', textinfo='percent+label')
        fig_qty.update_layout(showlegend=False, margin=dict(t=40, b=10, l=10, r=10))

        fig_uf = px.pie(df_pie_uf, values='UF', names='Entregable', title='Distribución por Valorización (UF)', hole=0.3)
        fig_uf.update_traces(textposition='inside', textinfo='percent+label')
        fig_uf.update_layout(showlegend=False, margin=dict(t=40, b=10, l=10, r=10))
        
        c_pie1, c_pie2 = st.columns(2)
        with c_pie1: st.plotly_chart(fig_qty, use_container_width=True, key=f"pie_qty_{target_week_id}")
        with c_pie2: st.plotly_chart(fig_uf, use_container_width=True, key=f"pie_uf_{target_week_id}")
        
        df_entregables['honorarios_estimados'] = df_entregables['honorarios_estimados'].apply(lambda x: f"{x:,.2f}")
        
        df_mostrar = df_entregables.rename(columns={
            'numero_caso': 'Caso', 
            'asegurado': 'Asegurado', 
            'accion': 'Tipo de Entregable', 
            'honorarios_estimados': 'Hon UF'
        })
        st.dataframe(df_mostrar, use_container_width=True, hide_index=True)
    else:
        st.info("Aún no se han ejecutado entregables de valor (Informes o Presentaciones) esta semana.")
    st.markdown("---")
    
    # --- 5. GESTIÓN ESTRATÉGICA DESTACADA (CON FECHAS) ---
    st.markdown('<div class="marco-gestion" style="border-left: 5px solid #ffc107;"><h4>🤝 Cuadrante 4: Gestión Estratégica Transversal</h4></div>', unsafe_allow_html=True)
    
    df_com_raw = df_realizados[(df_realizados['categoria'] == 'Gestión Comercial') & (~df_realizados['accion'].isin(['0', ' ', '', 0]))].copy()
    df_com = pd.DataFrame()
    if not df_com_raw.empty:
        df_com_raw['Fecha'] = pd.to_datetime(df_com_raw['fecha_ejecucion'], errors='coerce').dt.strftime('%d-%m-%Y').fillna('N/D')
        df_com = df_com_raw[['Fecha', 'Ajustador', 'accion']].rename(columns={'accion': 'Detalle de la Gestión'})
        df_com = df_com.drop_duplicates() 
    
    df_adm_raw = df_realizados[(df_realizados['categoria'] == 'Gestión Administrativa') & (~df_realizados['accion'].isin(['0', ' ', '', 0]))].copy()
    df_adm = pd.DataFrame()
    if not df_adm_raw.empty:
        df_adm_raw['Fecha'] = pd.to_datetime(df_adm_raw['fecha_ejecucion'], errors='coerce').dt.strftime('%d-%m-%Y').fillna('N/D')
        df_adm = df_adm_raw[['Fecha', 'Ajustador', 'accion']].rename(columns={'accion': 'Detalle / Comités'})
        df_adm = df_adm.drop_duplicates() 
    
    col_com, col_adm = st.columns(2)
    with col_com:
        st.markdown(f"<div style='background-color:#e6f2ff; padding:15px; border-radius:8px; border-left: 4px solid #004a99;'><h5 style='color:#004a99; margin-top:0;'>Reuniones Comerciales ({len(df_com)})</h5></div>", unsafe_allow_html=True)
        if not df_com.empty:
            st.dataframe(df_com, use_container_width=True, hide_index=True)
        else:
            st.caption("Sin gestiones comerciales válidas reportadas.")
    with col_adm:
        st.markdown(f"<div style='background-color:#f8f9fa; padding:15px; border-radius:8px; border-left: 4px solid #6c757d;'><h5 style='color:#6c757d; margin-top:0;'>Gestiones Administrativas ({len(df_adm)})</h5></div>", unsafe_allow_html=True)
        if not df_adm.empty:
            st.dataframe(df_adm, use_container_width=True, hide_index=True)
        else:
            st.caption("Sin gestiones administrativas válidas reportadas.")
            
    # --- GENERADORES EXPORTACIÓN DASHBOARD ---
    dash_wb = Workbook()
    ws_resumen = dash_wb.active
    ws_resumen.title = "Resumen Ejecutivo"
    
    ws_resumen.append(["Métrica Financiera", "UF"])
    ws_resumen.append(["Ingreso Efectivo Facturable (Cierres/Rechazos)", round(uf_facturables_caja, 2)])
    ws_resumen.append(["Valor Potencial Traccionado (WIP / Proceso)", round(uf_traccionadas_wip, 2)])
    ws_resumen.append([])
    
    ws_resumen.append(["KPI Operativo", "Valor (%)"])
    ws_resumen.append(["Adherencia al Plan", round(adherencia, 1)])
    ws_resumen.append(["Ratio de Esfuerzo Planificado", round(ratio_plan, 1)])
    
    excel_dash_buffer = io.BytesIO()
    dash_wb.save(excel_dash_buffer)

    # --- DOCUMENTO WORD - DISEÑO TIPO DASHBOARD ---
    dash_doc = Document()
    dash_doc.add_heading(f'📊 Dashboard Ejecutivo y Cumplimiento - {week_id_obj}', 0)
    
    # 1. Tabla Financiera VIP (Se mantiene grande en la portada)
    dash_doc.add_heading('💰 1. Valorización de Cartera y Facturación Proyectada', level=1)
    t_fin = dash_doc.add_table(rows=2, cols=2)
    t_fin.style = 'Table Grid'
    
    t_fin.rows[0].cells[0].text = "Ingreso Efectivo Facturable (Caja)"
    t_fin.rows[0].cells[1].text = "Valor Potencial Traccionado (WIP)"
    formatear_cabecera_tabla(t_fin, "003366")
    
    celda_caja = t_fin.rows[1].cells[0]
    celda_caja.text = f"{uf_facturables_caja:,.2f} UF"
    celda_caja.paragraphs[0].runs[0].font.size = Pt(22)
    celda_caja.paragraphs[0].runs[0].font.bold = True
    celda_caja.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 51, 102)
    celda_caja.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    celda_wip = t_fin.rows[1].cells[1]
    celda_wip.text = f"{uf_traccionadas_wip:,.2f} UF"
    celda_wip.paragraphs[0].runs[0].font.size = Pt(22)
    celda_wip.paragraphs[0].runs[0].font.bold = True
    celda_wip.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 51, 102)
    celda_wip.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    dash_doc.add_paragraph("")
    
    # 2. Tabla KPIs (MOTOR MATPLOTLIB OFFLINE)
    dash_doc.add_heading('⏱️ 2. Métricas de Cumplimiento y Carga Operativa', level=1)
    t_kpi = dash_doc.add_table(rows=1, cols=2)
    t_kpi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def generar_velocimetro_estatico(valor, titulo):
        import numpy as np
        import matplotlib.pyplot as plt
        import matplotlib.patches as patches
        valor_seguro = min(max(float(valor), 0.0), 100.0)
        fig, ax = plt.subplots(figsize=(5, 3.2))
        center, radius, width = (0, 0), 1.0, 0.3
        
        def draw_wedge(start, end, color):
            wedge = patches.Wedge(center, radius, start, end, width=width, facecolor=color, edgecolor='gray', linewidth=1)
            ax.add_patch(wedge)
            
        draw_wedge(90, 180, '#d9534f')
        draw_wedge(36, 90, '#f0ad4e')
        draw_wedge(0, 36, '#28a745')

        angle_progress = 180 - (valor_seguro * 1.8)
        theta = np.linspace(np.radians(180), np.radians(angle_progress), 100)
        r_line = radius - (width / 2)
        ax.plot(r_line * np.cos(theta), r_line * np.sin(theta), color='black', linewidth=8, solid_capstyle='round')

        for t in [0, 20, 40, 60, 80, 100]:
            ang = np.radians(180 - t * 1.8)
            ax.plot([radius * np.cos(ang), (radius + 0.05) * np.cos(ang)], [radius * np.sin(ang), (radius + 0.05) * np.sin(ang)], color='gray', lw=1.5)
            ax.text((radius + 0.15) * np.cos(ang), (radius + 0.15) * np.sin(ang), str(t), ha='center', va='center', fontsize=10, color='#333333')

        ax.text(0, 0.15, f"{valor_seguro:.1f}%", ha='center', va='center', fontsize=28, fontweight='bold')
        ax.text(0, 1.25, titulo, ha='center', va='center', fontsize=16, color='#003366', fontweight='bold')
        ax.set_xlim(-1.3, 1.3)
        ax.set_ylim(-0.1, 1.4)
        ax.axis('off')

        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', transparent=True)
        plt.close(fig)
        return buf.getvalue()

    try:
        img_adh_bytes = generar_velocimetro_estatico(adherencia, "Adherencia al Plan")
        celda_img_adh = t_kpi.rows[0].cells[0]
        celda_img_adh.paragraphs[0].add_run().add_picture(io.BytesIO(img_adh_bytes), width=Cm(7.5))
        celda_img_adh.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        img_pro_bytes = generar_velocimetro_estatico(ratio_plan, "Ratio Planificado")
        celda_img_pro = t_kpi.rows[0].cells[1]
        celda_img_pro.paragraphs[0].add_run().add_picture(io.BytesIO(img_pro_bytes), width=Cm(7.5))
        celda_img_pro.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        t_kpi.rows[0].cells[0].text = f"Adherencia: {adherencia:.1f}%"
        t_kpi.rows[0].cells[1].text = f"Ratio Planificado: {ratio_plan:.1f}%"
        
    # 3. Tabla MCL (CON SALTO DE PÁGINA Y FUENTE REDUCIDA)
    if not df_mcl_mostrar.empty:
        dash_doc.add_page_break()
        dash_doc.add_heading('🏆 Radar de Casos MCL', level=1)
        t_mcl = dash_doc.add_table(rows=1, cols=len(df_mcl_mostrar.columns))
        t_mcl.style = 'Table Grid'
        columnas_mcl = list(df_mcl_mostrar.columns)
        for i, col_name in enumerate(columnas_mcl):
            t_mcl.rows[0].cells[i].text = str(col_name)
            t_mcl.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        formatear_cabecera_tabla(t_mcl, "D9534F")
        
        for row_val in df_mcl_mostrar.values.tolist():
            row_cells = t_mcl.add_row().cells
            for i, val in enumerate(row_val):
                row_cells[i].text = str(val)
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(8.5)
    
    # 4. Tabla Entregables con Gráficos de Torta (CON SALTO DE PÁGINA Y FUENTE REDUCIDA)
    if not df_mostrar.empty:
        dash_doc.add_page_break()
        dash_doc.add_heading('🏭 3. Inventario de Producción (Entregables)', level=1)
        
        def generar_torta_estatica(labels, sizes, titulo):
            import matplotlib.pyplot as plt
            import io
            fig, ax = plt.subplots(figsize=(6, 5))
            wedges, texts, autotexts = ax.pie(
                sizes, 
                autopct='%1.1f%%', 
                startangle=90, 
                pctdistance=0.75,
                textprops=dict(color="w", weight="bold", fontsize=9),
                wedgeprops=dict(edgecolor='w', linewidth=1)
            )
            ax.legend(
                wedges, labels, 
                title="Entregables", 
                loc="upper center", 
                bbox_to_anchor=(0.5, -0.05), 
                fontsize=9
            )
            ax.set_title(titulo, fontsize=12, color='#003366', fontweight='bold', pad=15)
            buf = io.BytesIO()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', transparent=True)
            plt.close(fig)
            return buf.getvalue()

        try:
            t_pies = dash_doc.add_table(rows=1, cols=2)
            t_pies.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if df_pie_qty['Cantidad'].sum() > 0:
                img_qty = generar_torta_estatica(df_pie_qty['Entregable'].tolist(), df_pie_qty['Cantidad'].tolist(), "Cantidad de Entregables")
                c_qty = t_pies.rows[0].cells[0]
                c_qty.paragraphs[0].add_run().add_picture(io.BytesIO(img_qty), width=Cm(7.5))
                c_qty.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if df_pie_uf['UF'].sum() > 0:
                img_uf = generar_torta_estatica(df_pie_uf['Entregable'].tolist(), df_pie_uf['UF'].tolist(), "Valorización en UF")
                c_uf = t_pies.rows[0].cells[1]
                c_uf.paragraphs[0].add_run().add_picture(io.BytesIO(img_uf), width=Cm(7.5))
                c_uf.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            dash_doc.add_paragraph("")
        except Exception:
            pass

        t_ent = dash_doc.add_table(rows=1, cols=len(df_mostrar.columns))
        t_ent.style = 'Table Grid'
        for i, col_name in enumerate(df_mostrar.columns):
            t_ent.rows[0].cells[i].text = str(col_name)
            t_ent.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        formatear_cabecera_tabla(t_ent, "28A745")
        
        for row_val in df_mostrar.values.tolist():
            row_cells = t_ent.add_row().cells
            for i, val in enumerate(row_val):
                row_cells[i].text = str(val)
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(8.5)
        
    # 5. Tablas Comerciales y Admin (CON SALTO DE PÁGINA Y FUENTE REDUCIDA)
    if not df_com.empty or not df_adm.empty:
        dash_doc.add_page_break()
        dash_doc.add_heading('🤝 4. Gestión Estratégica Transversal', level=1)
        
        if not df_com.empty:
            dash_doc.add_heading('🔵 Gestiones Comerciales', level=2)
            t_com_w = dash_doc.add_table(rows=1, cols=3)
            t_com_w.style = 'Table Grid'
            encabezados_com = ["Fecha", "Ajustador", "Detalle de la Gestión"]
            for i, titulo in enumerate(encabezados_com):
                t_com_w.rows[0].cells[i].text = titulo
                t_com_w.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
            formatear_cabecera_tabla(t_com_w, "004A99") 
            
            for _, row in df_com.iterrows():
                row_cells = t_com_w.add_row().cells
                row_cells[0].text = str(row['Fecha'])
                row_cells[1].text = str(row['Ajustador'])
                row_cells[2].text = str(row['Detalle de la Gestión'])
                for cell in row_cells:
                    cell.paragraphs[0].runs[0].font.size = Pt(8.5)
            dash_doc.add_paragraph("")
        
        if not df_adm.empty:
            dash_doc.add_heading('⚪ Gestiones Administrativas', level=2)
            t_adm_w = dash_doc.add_table(rows=1, cols=3)
            t_adm_w.style = 'Table Grid'
            encabezados_adm = ["Fecha", "Ajustador", "Detalle / Comités"]
            for i, titulo in enumerate(encabezados_adm):
                t_adm_w.rows[0].cells[i].text = titulo
                t_adm_w.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
            formatear_cabecera_tabla(t_adm_w, "6C757D") 
            
            for _, row in df_adm.iterrows():
                row_cells = t_adm_w.add_row().cells
                row_cells[0].text = str(row['Fecha'])
                row_cells[1].text = str(row['Ajustador'])
                row_cells[2].text = str(row['Detalle / Comités'])
                for cell in row_cells:
                    cell.paragraphs[0].runs[0].font.size = Pt(8.5)

    word_dash_buffer = io.BytesIO()
    dash_doc.save(word_dash_buffer)
    
    # --- BOTONES DE DESCARGA EJECUTIVO ---
    st.markdown("---")
    st.markdown("### Opciones de Exportación Dashboard")
    col_d1, col_d2, col_d3 = st.columns(3)
    with col_d1:
        st.download_button(
            label="📥 DESCARGAR DASHBOARD (EXCEL)", 
            data=excel_dash_buffer.getvalue(), 
            file_name=f"Dashboard_Ejecutivo_{target_week_id}.xlsx", 
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    with col_d2:
        st.download_button(
            label="📥 DESCARGAR DASHBOARD (WORD)", 
            data=word_dash_buffer.getvalue(), 
            file_name=f"Resumen_Ejecutivo_{target_week_id}.docx", 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# ---------------------------------------------------------
# BLOQUE 4.4: VISTA - REPORTE OPERACIONAL DE EQUIPO
# VERSIÓN: 4.4.3 (Filtro Inteligente + ID Únicos para Gráficos)
# ---------------------------------------------------------
def renderizar_reporte_operacional(df_week, ajustadores_validos, target_week_id, week_id_obj):
    import io
    import pandas as pd
    import streamlit as st
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    st.markdown("### 📋 Radiografía Operacional por Ajustador")
    st.markdown("Análisis individual para reuniones de seguimiento: Embudo, Cumplimiento, Tareas No Programadas y Gestión.")

    op_doc = Document()
    op_doc.add_heading(f'📋 Reporte Operacional Detallado por Ajustador - {week_id_obj}', 0)

    if not ajustadores_validos:
        st.warning("No se pudo cargar la lista de ajustadores desde la Base Maestra.")
        return

    for i, ajustador in enumerate(ajustadores_validos):

        # --- SALTO DE PÁGINA: FICHA INDIVIDUAL POR AJUSTADOR ---
        if i > 0:
            op_doc.add_page_break()

        op_doc.add_heading(f"👤 {ajustador}", level=1)

        # Filtro flexible: ignora mayúsculas y espacios extra para emparejar bien la semana pasada
        df_aj = df_week[df_week['Ajustador'].astype(str).str.strip().str.lower() == str(ajustador).strip().lower()].copy()

        st.markdown(f"#### 👤 {ajustador}")

        if df_aj.empty:
            st.markdown("<h4 style='color:#fff; background-color:#d9534f; padding:15px; border-radius:5px; text-align:center;'>🚨 AJUSTADOR SIN PLAN</h4>", unsafe_allow_html=True)

            t_alerta = op_doc.add_table(rows=1, cols=1)
            celda_alerta = t_alerta.rows[0].cells[0]
            celda_alerta.text = "🚨 AJUSTADOR SIN PLAN REPORTADO EN ESTE PERIODO"
            formatear_cabecera_tabla(t_alerta, "D9534F")
            celda_alerta.paragraphs[0].runs[0].font.size = Pt(14)
            celda_alerta.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            op_doc.add_paragraph("")
        else:
            df_aj_realizado = df_aj[df_aj['estado_cumplimiento'] == 'Realizado']

            cond_fac_aj = df_aj_realizado['accion'].str.contains('Informe Final de Liquidación|Carta de Cobertura \(Rechazo\)', case=False, na=False)
            uf_caja_aj = df_aj_realizado[cond_fac_aj]['honorarios_estimados'].sum()

            cond_wip_aj = (df_aj_realizado['categoria'] == 'Operativa') & (~cond_fac_aj)
            uf_wip_aj = df_aj_realizado[cond_wip_aj]['honorarios_estimados'].sum()

            t_prog = len(df_aj[df_aj['tipo_actividad'] == 'Programada'])
            t_prog_hechas = len(df_aj_realizado[df_aj_realizado['tipo_actividad'] == 'Programada'])
            t_np = len(df_aj_realizado[df_aj_realizado['tipo_actividad'] == 'Actividad Adicional'])

            if t_prog > 0:
                adh_aj = (t_prog_hechas / t_prog) * 100
            else:
                adh_aj = 0

            total_aj = t_prog_hechas + t_np

            if total_aj > 0:
                rat_prog = (t_prog_hechas / total_aj) * 100
                rat_np = (t_np / total_aj) * 100
            else:
                rat_prog = 0
                rat_np = 0

            colA, colB = st.columns(2)
            with colA:
                st.metric("Facturación Lograda (UF)", f"{uf_caja_aj:,.2f}")
            with colB:
                st.metric("Esfuerzo en Proceso (UF)", f"{uf_wip_aj:,.2f}")

            fig_adh_op = crear_velocimetro(adh_aj, "Adherencia")
            fig_pro_op = crear_velocimetro(rat_prog, "Proactivo")

            c_op1, c_op2 = st.columns(2)
            # SOLUCIÓN AL ERROR: Creación de Keys Únicos concatenando nombre, semana e índice
            str_aj = str(ajustador).replace(' ', '_').replace('.', '')
            with c_op1:
                st.plotly_chart(fig_adh_op, use_container_width=True, key=f"op_adh_{str_aj}_{target_week_id}_{i}")
            with c_op2:
                st.plotly_chart(fig_pro_op, use_container_width=True, key=f"op_pro_{str_aj}_{target_week_id}_{i}")

            t_fin_op = op_doc.add_table(rows=2, cols=2)
            t_fin_op.style = 'Table Grid'
            t_fin_op.rows[0].cells[0].text = "💰 Facturación Lograda (Caja)"
            t_fin_op.rows[0].cells[1].text = "📈 Esfuerzo en Proceso (WIP)"
            formatear_cabecera_tabla(t_fin_op, "003366")

            celda_caja_op = t_fin_op.rows[1].cells[0]
            celda_caja_op.text = f"{uf_caja_aj:,.2f} UF"
            celda_caja_op.paragraphs[0].runs[0].font.size = Pt(20)
            celda_caja_op.paragraphs[0].runs[0].font.bold = True
            celda_caja_op.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 51, 102)
            celda_caja_op.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            celda_wip_op = t_fin_op.rows[1].cells[1]
            celda_wip_op.text = f"{uf_wip_aj:,.2f} UF"
            celda_wip_op.paragraphs[0].runs[0].font.size = Pt(20)
            celda_wip_op.paragraphs[0].runs[0].font.bold = True
            celda_wip_op.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 51, 102)
            celda_wip_op.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            op_doc.add_paragraph("")

            # MOTOR MATPLOTLIB OFFLINE PARA REPORTE OPERACIONAL (Exportación Word)
            t_kpi_op = op_doc.add_table(rows=1, cols=2)
            t_kpi_op.alignment = WD_ALIGN_PARAGRAPH.CENTER

            def generar_velocimetro_estatico(valor, titulo):
                import numpy as np
                import matplotlib.pyplot as plt
                import matplotlib.patches as patches
                valor_seguro = min(max(float(valor), 0.0), 100.0)
                fig, ax = plt.subplots(figsize=(5, 3.2))
                center, radius, width = (0, 0), 1.0, 0.3

                def draw_wedge(start, end, color):
                    wedge = patches.Wedge(center, radius, start, end, width=width, facecolor=color, edgecolor='gray', linewidth=1)
                    ax.add_patch(wedge)

                draw_wedge(90, 180, '#d9534f')
                draw_wedge(36, 90, '#f0ad4e')
                draw_wedge(0, 36, '#28a745')

                angle_progress = 180 - (valor_seguro * 1.8)
                theta = np.linspace(np.radians(180), np.radians(angle_progress), 100)
                r_line = radius - (width / 2)
                ax.plot(r_line * np.cos(theta), r_line * np.sin(theta), color='black', linewidth=8, solid_capstyle='round')

                for t in [0, 20, 40, 60, 80, 100]:
                    ang = np.radians(180 - t * 1.8)
                    ax.plot([radius * np.cos(ang), (radius + 0.05) * np.cos(ang)], [radius * np.sin(ang), (radius + 0.05) * np.sin(ang)], color='gray', lw=1.5)
                    ax.text((radius + 0.15) * np.cos(ang), (radius + 0.15) * np.sin(ang), str(t), ha='center', va='center', fontsize=10, color='#333333')

                ax.text(0, 0.15, f"{valor_seguro:.1f}%", ha='center', va='center', fontsize=28, fontweight='bold')
                ax.text(0, 1.25, titulo, ha='center', va='center', fontsize=16, color='#003366', fontweight='bold')
                ax.set_xlim(-1.3, 1.3)
                ax.set_ylim(-0.1, 1.4)
                ax.axis('off')

                buf = io.BytesIO()
                plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', transparent=True)
                plt.close(fig)
                return buf.getvalue()

            try:
                img_adh_op_bytes = generar_velocimetro_estatico(adh_aj, "Adherencia")
                celda_img_adh_op = t_kpi_op.rows[0].cells[0]
                celda_img_adh_op.paragraphs[0].add_run().add_picture(io.BytesIO(img_adh_op_bytes), width=Cm(7.5))
                celda_img_adh_op.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                img_pro_op_bytes = generar_velocimetro_estatico(rat_prog, "Proactivo")
                celda_img_pro_op = t_kpi_op.rows[0].cells[1]
                celda_img_pro_op.paragraphs[0].add_run().add_picture(io.BytesIO(img_pro_op_bytes), width=Cm(7.5))
                celda_img_pro_op.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                celda_img_adh_op = t_kpi_op.rows[0].cells[0]
                celda_img_adh_op.text = f"Adherencia: {adh_aj:.1f}%"
                celda_img_pro_op = t_kpi_op.rows[0].cells[1]
                celda_img_pro_op.text = f"Proactivo: {rat_prog:.1f}%"

            op_doc.add_paragraph("")

            with st.expander(f"Ver desglose operativo detallado"):

                df_prog_view = df_aj_realizado[df_aj_realizado['tipo_actividad'] == 'Programada'][['numero_caso', 'asegurado', 'accion', 'honorarios_estimados']]
                if not df_prog_view.empty:
                    st.markdown("**✅ Trabajo Programado Realizado:**")
                    st.dataframe(df_prog_view, use_container_width=True, hide_index=True)

                    op_doc.add_heading('✅ Trabajo Programado Realizado:', level=2)
                    t_prog_w = op_doc.add_table(rows=1, cols=4)
                    t_prog_w.style = 'Table Grid'

                    t_prog_w.rows[0].cells[0].text = "Caso"
                    t_prog_w.rows[0].cells[1].text = "Asegurado"
                    t_prog_w.rows[0].cells[2].text = "Acción"
                    t_prog_w.rows[0].cells[3].text = "Hon UF"

                    formatear_cabecera_tabla(t_prog_w, "28A745") 

                    for r_v in df_prog_view.values.tolist():
                        r_c = t_prog_w.add_row().cells
                        r_c[0].text = str(r_v[0])
                        r_c[1].text = str(r_v[1])
                        r_c[2].text = str(r_v[2])
                        try:
                            r_c[3].text = f"{float(r_v[3]):,.2f}"
                        except:
                            r_c[3].text = "0.00"
                    op_doc.add_paragraph("")

                if t_np > 0:
                    st.markdown("**🔴 Detalle de Tareas No Programadas:**")
                    df_np_view = df_aj_realizado[df_aj_realizado['tipo_actividad'] == 'Actividad Adicional'][['numero_caso', 'asegurado', 'accion', 'fecha_ejecucion']]
                    st.dataframe(df_np_view, use_container_width=True, hide_index=True)

                    op_doc.add_heading('🔴 Tareas No Programadas:', level=2)
                    t_ur = op_doc.add_table(rows=1, cols=4)
                    t_ur.style = 'Table Grid'

                    for j_col, col in enumerate(df_np_view.columns): 
                        t_ur.rows[0].cells[j_col].text = str(col).capitalize()

                    formatear_cabecera_tabla(t_ur, "D9534F") 

                    for r_v in df_np_view.values.tolist():
                        r_c = t_ur.add_row().cells
                        for j_col, v in enumerate(r_v): 
                            r_c[j_col].text = str(v)
                    op_doc.add_paragraph("")

                cond_estr_aj = (df_aj_realizado['categoria'].isin(['Gestión Comercial', 'Gestión Administrativa'])) & (~df_aj_realizado['accion'].isin(['0', ' ', '', 0]))
                df_estr = df_aj_realizado[cond_estr_aj]

                if not df_estr.empty:
                    st.markdown("**🔵 Gestiones Comerciales y Administrativas:**")
                    st.dataframe(df_estr[['categoria', 'accion']], use_container_width=True, hide_index=True)

                    op_doc.add_heading('🔵 Gestión Estratégica Transversal:', level=2)
                    t_est_w = op_doc.add_table(rows=1, cols=2)
                    t_est_w.style = 'Table Grid'

                    t_est_w.rows[0].cells[0].text = "Categoría"
                    t_est_w.rows[0].cells[1].text = "Acción / Detalle"

                    formatear_cabecera_tabla(t_est_w, "17A2B8") 

                    for _, r in df_estr.iterrows():
                        r_c = t_est_w.add_row().cells
                        r_c[0].text = str(r['categoria'])
                        r_c[1].text = str(r['accion'])

            op_doc.add_paragraph("")
        st.markdown("---")

    # --- BOTONES DE DESCARGA OPERACIONAL ---
    st.markdown("### Exportar Reporte Operacional")
    word_op_buffer = io.BytesIO()
    op_doc.save(word_op_buffer)
    st.download_button(
        label="📥 DESCARGAR REPORTE OPERACIONAL (WORD)", 
        data=word_op_buffer.getvalue(), 
        file_name=f"Reporte_Operacional_{target_week_id}.docx", 
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
        type="primary"
    )


# ---------------------------------------------------------
# BLOQUE 4.5: VISTA - CARTA GANTT OPERATIVA
# VERSIÓN: 4.5.16 (Corrección Quirúrgica de NameError df_aj)
# ---------------------------------------------------------
def renderizar_carta_gantt(df_week, df_raw, dias_semana_target, target_week_id, week_id_obj):
    import io
    import os
    import json
    import pandas as pd
    import streamlit as st
    import plotly.express as px
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from docx import Document
    from docx.shared import RGBColor, Cm, Pt
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from datetime import datetime

    if df_week.empty:
        st.info("No hay tareas operativas planificadas aún para esta semana.")
        return
        
    if 'categoria' in df_week.columns:
        df_operativa = df_week[df_week['categoria'] == 'Operativa'].copy()
    else:
        df_operativa = df_week.copy()
    
    if not df_operativa.empty:
        df_operativa['fecha_compromiso'] = pd.to_datetime(df_operativa['fecha_compromiso'], errors='coerce').dt.date
        
        # Filtro estricto: Expulsar cualquier tarea que caiga en Sábado (5) o Domingo (6)
        df_operativa = df_operativa[pd.to_datetime(df_operativa['fecha_compromiso']).apply(lambda x: x.weekday() < 5)]
        
        if df_operativa.empty:
            st.info("No hay tareas operativas planificadas en días hábiles para esta semana.")
            return

        if 'estado_proyectado' not in df_operativa.columns:
            df_operativa['estado_proyectado'] = 'N/D'

        # Limpieza crucial para que la suma aritmética no falle al leer strings
        df_operativa['honorarios_estimados'] = pd.to_numeric(df_operativa['honorarios_estimados'], errors='coerce').fillna(0.0)

        # --- MAPEO DE DIVISIONES DESDE LA BASE MAESTRA ---
        dict_divisiones = {}
        filepath = os.path.join(PERSISTENCE_DIR, "BASE_MAESTRA.json")
        if os.path.exists(filepath):
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    datos = json.load(f)
                    df_maestro = pd.DataFrame(datos['data'])
                    if not df_maestro.empty:
                        col_div = df_maestro.columns[3] # Columna D
                        col_aj = 'Ajustador senior' if 'Ajustador senior' in df_maestro.columns else df_maestro.columns[9]
                        for _, row in df_maestro.iterrows():
                            aj_name = str(row[col_aj]).strip()
                            div_name = str(row[col_div]).strip()
                            if aj_name and div_name:
                                dict_divisiones[aj_name] = div_name
            except Exception:
                pass
                
        df_operativa['Division'] = df_operativa['Ajustador'].apply(lambda x: dict_divisiones.get(str(x).strip(), 'Sin División Asignada'))

        # --- CÁLCULO DE TOTALES (Estricto por Caso y Determinación de Origen) ---
        uf_ifl_total = 0.0
        uf_wip_total = 0.0
        uf_mcl_total = 0.0
        
        # Sub-baldes para gráficos de torta
        ifl_prog, ifl_noprog = 0.0, 0.0
        wip_prog, wip_noprog = 0.0, 0.0
        
        resumen_ajustadores = {}
        casos_agrupados = df_operativa.groupby(['Division', 'Ajustador', 'numero_caso'])
        
        for (division, ajustador, caso), group in casos_agrupados:
            try:
                uf_val = float(group['honorarios_estimados'].max())
            except:
                uf_val = 0.0
                
            acciones_del_caso = " ".join(group['accion'].astype(str)).lower()
            es_ifl = 'informe final de liquidación' in acciones_del_caso or 'carta de cobertura (rechazo)' in acciones_del_caso
            
            tramo_str = str(group['tramo_uf'].iloc[0]).lower()
            es_mcl = 'mcl' in tramo_str or '> 5' in tramo_str
            
            # Lógica de Origen: Si al menos 1 tarea del caso es Programada, se cuenta como Programado
            tipos_act = group['tipo_actividad'].astype(str).tolist()
            es_programado = any(t == 'Programada' for t in tipos_act)

            if es_ifl:
                uf_ifl_total += uf_val
                if es_programado: ifl_prog += uf_val
                else: ifl_noprog += uf_val
            else:
                uf_wip_total += uf_val
                if es_programado: wip_prog += uf_val
                else: wip_noprog += uf_val
                
            if es_mcl:
                uf_mcl_total += uf_val

            if ajustador not in resumen_ajustadores:
                resumen_ajustadores[ajustador] = {'Division': division, 'IFL': 0.0, 'WIP': 0.0, 'MCL': 0.0, 'Total': 0.0}

            if es_ifl: resumen_ajustadores[ajustador]['IFL'] += uf_val
            else: resumen_ajustadores[ajustador]['WIP'] += uf_val
            if es_mcl: resumen_ajustadores[ajustador]['MCL'] += uf_val
            resumen_ajustadores[ajustador]['Total'] += uf_val

        # --- VISUALIZACIÓN UI (STREAMLIT) ---
        st.subheader("🛠️ Gantt Operativo (Línea de tiempo de la gerencia)")
        
        c_ui1, c_ui2 = st.columns(2)
        with c_ui1:
            st.markdown(f"""
            <div style='background-color: #28a745; padding: 25px; border-radius: 8px; text-align: center; margin-bottom: 20px;'>
                <div style='color: white; font-size: 18px; font-weight: bold; margin-bottom: 5px;'>IFL Comprometidos Globales</div>
                <div style='color: white; font-size: 30px; font-weight: bold;'>{uf_ifl_total:,.2f} UF</div>
            </div>
            """, unsafe_allow_html=True)
        with c_ui2:
            st.markdown(f"""
            <div style='background-color: #004A99; padding: 25px; border-radius: 8px; text-align: center; margin-bottom: 20px;'>
                <div style='color: white; font-size: 18px; font-weight: bold; margin-bottom: 5px;'>WIP Comprometidos Globales</div>
                <div style='color: white; font-size: 30px; font-weight: bold;'>{uf_wip_total:,.2f} UF</div>
            </div>
            """, unsafe_allow_html=True)

        # Gráficos de Torta UI
        col_pie1, col_pie2 = st.columns(2)
        
        df_pie_ifl = pd.DataFrame([['Plan Oficial', ifl_prog], ['Urgencia / No Prog', ifl_noprog]], columns=['Origen', 'UF'])
        df_pie_wip = pd.DataFrame([['Plan Oficial', wip_prog], ['Urgencia / No Prog', wip_noprog]], columns=['Origen', 'UF'])
        
        colores_torta = {'Plan Oficial': '#217346', 'Urgencia / No Prog': '#5BC0DE'}
        
        with col_pie1:
            if uf_ifl_total > 0:
                fig_ifl = px.pie(df_pie_ifl, values='UF', names='Origen', title='Composición IFL (Programado vs Urgencias)', hole=0.3, color='Origen', color_discrete_map=colores_torta)
                fig_ifl.update_traces(textposition='inside', textinfo='percent+label')
                fig_ifl.update_layout(showlegend=False, margin=dict(t=40, b=10, l=10, r=10))
                st.plotly_chart(fig_ifl, use_container_width=True, key=f"gantt_pie_ifl_{target_week_id}")
            
        with col_pie2:
            if uf_wip_total > 0:
                fig_wip = px.pie(df_pie_wip, values='UF', names='Origen', title='Composición WIP (Programado vs Urgencias)', hole=0.3, color='Origen', color_discrete_map=colores_torta)
                fig_wip.update_traces(textposition='inside', textinfo='percent+label')
                fig_wip.update_layout(showlegend=False, margin=dict(t=40, b=10, l=10, r=10))
                st.plotly_chart(fig_wip, use_container_width=True, key=f"gantt_pie_wip_{target_week_id}")

        df_gantt_visual = df_operativa.sort_values(by=['Division', 'Ajustador', 'numero_caso', 'fecha_compromiso'])
        df_gantt_ui = df_gantt_visual.pivot_table(
            index=['Division', 'Ajustador', 'numero_caso', 'Nick Name', 'asegurado'], 
            columns='fecha_compromiso', 
            values='accion', 
            aggfunc=lambda x: ' | '.join(x)
        ).fillna('')
        st.dataframe(df_gantt_ui, use_container_width=True)

        dias_habiles_target = [d for d in dias_semana_target if d.weekday() < 5]

        # --- EXPORTACIÓN EXCEL GANTT ---
        wb = Workbook()
        ws = wb.active
        ws.title = "Plan Semanal Gantt"
        
        encabezados = ["División", "Ajustador", "Caso", "Nick Name", "Asegurado", "Acción y Entregable"]
        for f in dias_habiles_target:
            encabezados.append(f.strftime('%A %d-%m'))
        encabezados.append("Hon UF")
        ws.append(encabezados)
        
        grouped_gantt_xl = df_operativa.groupby(['Division', 'Ajustador', 'numero_caso', 'Nick Name', 'asegurado', 'accion'])
        for name, group in grouped_gantt_xl:
            row = list(name)
            for f in dias_habiles_target:
                if f in group['fecha_compromiso'].values:
                    row.append("X")
                else:
                    row.append("")
            try:
                total_honorarios = float(group['honorarios_estimados'].max())
            except:
                total_honorarios = 0.0
            row.append(round(total_honorarios, 2))
            ws.append(row)
            
        header_fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = Font(color="FFFFFF", bold=True)
            cell.alignment = Alignment(horizontal="center", vertical="center")
        excel_buffer = io.BytesIO()
        wb.save(excel_buffer)
        
        # --- EXPORTACIÓN WORD GANTT (PORTADA CON GRÁFICOS MATPLOTLIB) ---
        doc = Document()
        section = doc.sections[-1]
        
        new_width = section.page_height
        new_height = section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

        doc.add_heading(f'📊 Reporte Consolidado de Planificación Semanal - {week_id_obj}', 0)
        
        # Carátula Global
        t_caratula = doc.add_table(rows=1, cols=2)
        t_caratula.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        c_ifl = t_caratula.rows[0].cells[0]
        shd_ifl = parse_xml(r'<w:shd {} w:fill="28A745"/>'.format(nsdecls('w')))
        c_ifl._tc.get_or_add_tcPr().append(shd_ifl)
        p_ifl = c_ifl.paragraphs[0]
        p_ifl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_ifl_title = p_ifl.add_run("IFL Comprometidos de la Semana\n")
        r_ifl_title.font.color.rgb = RGBColor(255, 255, 255)
        r_ifl_title.font.bold = True
        r_ifl_val = p_ifl.add_run(f"{uf_ifl_total:,.2f} UF")
        r_ifl_val.font.color.rgb = RGBColor(255, 255, 255)
        r_ifl_val.font.size = Pt(24)
        r_ifl_val.font.bold = True
        
        c_wip = t_caratula.rows[0].cells[1]
        shd_wip = parse_xml(r'<w:shd {} w:fill="004A99"/>'.format(nsdecls('w')))
        c_wip._tc.get_or_add_tcPr().append(shd_wip)
        p_wip = c_wip.paragraphs[0]
        p_wip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_wip_title = p_wip.add_run("WIP Comprometidos de la Semana\n")
        r_wip_title.font.color.rgb = RGBColor(255, 255, 255)
        r_wip_title.font.bold = True
        r_wip_val = p_wip.add_run(f"{uf_wip_total:,.2f} UF")
        r_wip_val.font.color.rgb = RGBColor(255, 255, 255)
        r_wip_val.font.size = Pt(24)
        r_wip_val.font.bold = True

        doc.add_paragraph("")
        
        # --- INYECCIÓN DE GRÁFICOS DE TORTA ESTÁTICOS EN WORD ---
        def generar_torta_estatica(sizes, titulo):
            import matplotlib.pyplot as plt
            import io
            fig, ax = plt.subplots(figsize=(4.5, 3.2))
            # Tamaños: [Programado, No Programado]
            wedges, texts, autotexts = ax.pie(
                sizes, 
                autopct='%1.1f%%', 
                startangle=90, 
                pctdistance=0.75,
                colors=['#217346', '#5BC0DE'], # Verde, Celeste
                textprops=dict(color="w", weight="bold", fontsize=9),
                wedgeprops=dict(edgecolor='w', linewidth=1)
            )
            ax.legend(
                wedges, ['Plan Oficial', 'Urgencias'], 
                loc="upper center", 
                bbox_to_anchor=(0.5, -0.05), 
                fontsize=9
            )
            ax.set_title(titulo, fontsize=11, color='#003366', fontweight='bold', pad=15)
            buf = io.BytesIO()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', transparent=True)
            plt.close(fig)
            return buf.getvalue()
            
        try:
            t_pies = doc.add_table(rows=1, cols=2)
            t_pies.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if uf_ifl_total > 0:
                img_ifl = generar_torta_estatica([ifl_prog, ifl_noprog], "Origen de IFL (UF)")
                c_pie_ifl = t_pies.rows[0].cells[0]
                c_pie_ifl.paragraphs[0].add_run().add_picture(io.BytesIO(img_ifl), width=Cm(8))
                c_pie_ifl.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            if uf_wip_total > 0:
                img_wip = generar_torta_estatica([wip_prog, wip_noprog], "Origen de WIP (UF)")
                c_pie_wip = t_pies.rows[0].cells[1]
                c_pie_wip.paragraphs[0].add_run().add_picture(io.BytesIO(img_wip), width=Cm(8))
                c_pie_wip.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            doc.add_paragraph("")
        except Exception:
            pass
        
        # --- TABLA DE RESUMEN POR AJUSTADOR ---
        doc.add_heading("📋 Resumen de Carga de Trabajo Operativo", level=2)
        t_resumen = doc.add_table(rows=1, cols=6)
        t_resumen.style = 'Table Grid'
        encabezados_resumen = ["División", "Ajustador", "IFL (UF)", "Ajustes/WIP (UF)", "Total Cartera (UF)", ">> De los cuales MCL"]
        
        for i, title in enumerate(encabezados_resumen):
            t_resumen.rows[0].cells[i].text = title
            t_resumen.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
            t_resumen.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            t_resumen.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            shading_elm = parse_xml(r'<w:shd {} w:fill="004A99"/>'.format(nsdecls('w')))
            t_resumen.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_elm)

        resumen_ordenado = sorted(resumen_ajustadores.items(), key=lambda x: (x[1]['Division'], x[0]))

        for aj, data in resumen_ordenado:
            row_cells = t_resumen.add_row().cells
            row_cells[0].text = data['Division']
            row_cells[1].text = aj
            row_cells[2].text = f"{data['IFL']:,.2f}"
            row_cells[3].text = f"{data['WIP']:,.2f}"
            row_cells[4].text = f"{round((data['Total']), 2):,.2f}"
            row_cells[5].text = f"{data['MCL']:,.2f}"
            
            for i in range(2, 6):
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        row_tot = t_resumen.add_row().cells
        row_tot[0].text = "TOTAL GERENCIA"
        row_tot[1].text = ""
        row_tot[2].text = f"{uf_ifl_total:,.2f}"
        row_tot[3].text = f"{uf_wip_total:,.2f}"
        row_tot[4].text = f"{round((uf_ifl_total + uf_wip_total), 2):,.2f}"
        row_tot[5].text = f"{uf_mcl_total:,.2f}"
        
        for i in range(6):
            if row_tot[i].paragraphs[0].runs:
                row_tot[i].paragraphs[0].runs[0].font.bold = True
            else:
                row_tot[i].paragraphs[0].add_run("").bold = True
                
            if i >= 2:
                row_tot[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            shading_elm = parse_xml(r'<w:shd {} w:fill="E6F2FF"/>'.format(nsdecls('w')))
            row_tot[i]._tc.get_or_add_tcPr().append(shading_elm)

        doc.add_paragraph("")
        
        # --- LEYENDA DE COLORES EN WORD ---
        leyenda = doc.add_paragraph()
        run_g = leyenda.add_run("🟢 Verde: ")
        run_g.bold = True
        leyenda.add_run("A tiempo  |  ")
        run_y = leyenda.add_run("🟡 Amarillo: ")
        run_y.bold = True
        leyenda.add_run("Con atraso  |  ")
        run_c = leyenda.add_run("🔵 Celeste: ")
        run_c.bold = True
        leyenda.add_run("Programado fuera del Plan Semanal  |  ")
        run_r = leyenda.add_run("🔴 Rojo con X: ")
        run_r.bold = True
        leyenda.add_run("No ejecutado  |  ")
        run_p = leyenda.add_run("🟩 Verde vacío: ")
        run_p.bold = True
        leyenda.add_run("Programado (Futuro)")
        leyenda.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        hoy = datetime.now().date()
        headers_word = ["Caso", "Nick Name", "Asegurado", "Acción/Entregable", "L", "M", "X", "J", "V", "Hon UF"]

        # --- ITERACIÓN POR DIVISIÓN Y POR AJUSTADOR ---
        divisiones = sorted(df_operativa['Division'].unique())
        
        for div in divisiones:
            df_div = df_operativa[df_operativa['Division'] == div]
            ajustadores_div = sorted(df_div['Ajustador'].unique())
            
            for aj in ajustadores_div:
                doc.add_page_break()
                doc.add_heading(f"{div} | Ajustador: {aj}", level=2)
                
                # ---> CORRECCIÓN QUIRÚRGICA: DEFINIR df_aj AL INICIO DEL CICLO <---
                df_aj = df_div[df_div['Ajustador'] == aj]
                
                # Rescate de los datos ya calculados del diccionario central
                aj_data = resumen_ajustadores.get(aj, {'IFL': 0.0, 'WIP': 0.0})
                aj_ifl = aj_data['IFL']
                aj_wip = aj_data['WIP']

                t_bar = doc.add_table(rows=1, cols=1)
                c_bar = t_bar.rows[0].cells[0]
                shd_bar = parse_xml(r'<w:shd {} w:fill="004A99"/>'.format(nsdecls('w')))
                c_bar._tc.get_or_add_tcPr().append(shd_bar)
                p_bar = c_bar.paragraphs[0]
                p_bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_bar = p_bar.add_run(f"RESUMEN SEMANAL  ||  IFL: {aj_ifl:,.2f} UF  |  WIP: {aj_wip:,.2f} UF")
                r_bar.font.color.rgb = RGBColor(255, 255, 255)
                r_bar.font.bold = True
                r_bar.font.size = Pt(12)
                doc.add_paragraph("")

                table = doc.add_table(rows=1, cols=len(headers_word))
                table.style = 'Table Grid'
                tr = table.rows[0]._tr
                trPr = tr.get_or_add_trPr()
                tblHeader = parse_xml(r'<w:tblHeader {} w:val="true"/>'.format(nsdecls('w')))
                trPr.append(tblHeader)
                
                for i, title in enumerate(headers_word):
                    table.rows[0].cells[i].text = title
                    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
                    table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    shading_elm = parse_xml(r'<w:shd {} w:fill="003366"/>'.format(nsdecls('w')))
                    table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_elm)

                caso_previo = ""
                grouped_aj = df_aj.groupby(['numero_caso', 'Nick Name', 'asegurado', 'accion'])
                
                for name, group in grouped_aj:
                    caso_actual, nickname, asegurado, accion = name
                    row_cells = table.add_row().cells
                    
                    if caso_actual == caso_previo:
                        row_cells[0].text = ""
                        row_cells[1].text = ""
                        row_cells[2].text = ""
                    else:
                        row_cells[0].text = str(caso_actual)
                        row_cells[1].text = str(nickname)
                        row_cells[2].text = str(asegurado)
                        
                    caso_previo = caso_actual
                    row_cells[3].text = str(accion)
                    
                    for i_date, f_date in enumerate(dias_habiles_target):
                        col_idx = 4 + i_date
                        match = group[group['fecha_compromiso'] == f_date]
                        
                        if not match.empty:
                            row_data = match.iloc[0]
                            estado_cumplimiento = str(row_data.get('estado_cumplimiento', ''))
                            fecha_ejec_str = str(row_data.get('fecha_ejecucion', ''))
                            tipo_actividad_gantt = str(row_data.get('tipo_actividad', ''))
                            
                            es_pasado = f_date < hoy
                            es_adicional = (tipo_actividad_gantt == 'Actividad Adicional')
                            
                            if es_pasado:
                                if estado_cumplimiento == 'Realizado':
                                    row_cells[col_idx].text = "✔"
                                    es_atrasada = False
                                    if pd.notna(fecha_ejec_str) and str(fecha_ejec_str).strip() not in ['', 'nan']:
                                        try:
                                            fecha_ejec_dt = pd.to_datetime(fecha_ejec_str).date()
                                            if fecha_ejec_dt > f_date: es_atrasada = True
                                        except: pass
                                        
                                    if es_adicional:
                                        shading_elm = parse_xml(r'<w:shd {} w:fill="5BC0DE"/>'.format(nsdecls('w')))
                                    elif es_atrasada:
                                        shading_elm = parse_xml(r'<w:shd {} w:fill="F0AD4E"/>'.format(nsdecls('w')))
                                    else:
                                        shading_elm = parse_xml(r'<w:shd {} w:fill="217346"/>'.format(nsdecls('w')))
                                else:
                                    row_cells[col_idx].text = "X"
                                    shading_elm = parse_xml(r'<w:shd {} w:fill="D9534F"/>'.format(nsdecls('w')))
                            else:
                                if estado_cumplimiento == 'Realizado':
                                    row_cells[col_idx].text = "✔"
                                    if es_adicional:
                                        shading_elm = parse_xml(r'<w:shd {} w:fill="5BC0DE"/>'.format(nsdecls('w')))
                                    else:
                                        shading_elm = parse_xml(r'<w:shd {} w:fill="217346"/>'.format(nsdecls('w')))
                                else:
                                    row_cells[col_idx].text = ""
                                    if es_adicional:
                                        shading_elm = parse_xml(r'<w:shd {} w:fill="5BC0DE"/>'.format(nsdecls('w')))
                                    else:
                                        shading_elm = parse_xml(r'<w:shd {} w:fill="217346"/>'.format(nsdecls('w')))
                                    
                            row_cells[col_idx]._tc.get_or_add_tcPr().append(shading_elm)
                        else:
                            row_cells[col_idx].text = ""
                    
                    try:
                        val_uf = float(group['honorarios_estimados'].max())
                        if val_uf > 0: row_cells[9].text = f"{val_uf:,.2f}"
                        else: row_cells[9].text = "-"
                    except: 
                        row_cells[9].text = "-"

                    for i, cell in enumerate(row_cells):
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs: 
                                run.font.size = Pt(7.5)
                            if i >= 4: 
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        word_buffer = io.BytesIO()
        doc.save(word_buffer)

        st.markdown("---")
        st.markdown("### Opciones de Exportación Gantt Corporativa")
        col1, col2, col3 = st.columns(3)
        with col1: 
            st.download_button(
                label="📥 DESCARGAR GANTT (EXCEL)", 
                data=excel_buffer.getvalue(), 
                file_name=f"Gantt_Planificacion_{target_week_id}.xlsx"
            )
        with col2: 
            st.download_button(
                label="📥 DESCARGAR REPORTE (WORD)", 
                data=word_buffer.getvalue(), 
                file_name=f"Reporte_Planificacion_{target_week_id}.docx"
            )
        with col3: 
            st.download_button(
                label="📥 DESCARGAR DATA BRUTA (CSV)", 
                data=df_raw.to_csv(index=False).encode('utf-8-sig'), 
                file_name=f"Data_Bruta_{target_week_id}.csv"
            )
    else:
        st.info("No hay tareas operativas (casos) planificadas aún para esta semana.")

# ---------------------------------------------------------
# BLOQUE 4.0: ORQUESTADOR PRINCIPAL DE LA VISTA
# VERSIÓN: 4.0.1 (Inclusión de Histórico Retroactivo - Semana Pasada)
# ---------------------------------------------------------
def vista_reportes():
    import streamlit as st
    from datetime import datetime, timedelta
    
    st.title("📊 Tablero de Control y Planificación")
    st.markdown("Visión gerencial del rendimiento financiero, cumplimiento operativo y línea de tiempo de la división.")
    
    col_radio, col_btn = st.columns([2, 1])
    with col_radio:
        week_id_obj = st.radio(
            "Seleccione la semana a reportar:", 
            ["Semana Pasada", "Semana Actual", "Próxima Semana"], 
            index=1,
            horizontal=True
        )
    with col_btn:
        st.markdown("<br>", unsafe_allow_html=True)
        forzar_sync = st.button("🔄 Sincronizar Nube ahora", type="primary", use_container_width=True)

    # Lógica de asignación de desplazamiento para cálculo de fechas
    if week_id_obj == "Semana Pasada":
        offset = -1
    elif week_id_obj == "Semana Actual":
        offset = 0
    else:
        offset = 1
        
    hoy = datetime.now()
    target_date = hoy + timedelta(weeks=offset)
    lunes = target_date - timedelta(days=target_date.weekday())
    dias_semana_target = []
    for i in range(7):
        dia_calculado = (lunes + timedelta(days=i)).date()
        dias_semana_target.append(dia_calculado)
        
    target_week_id = get_week_identifier(offset)

    # 4.2 Llamada al motor de datos
    df_week, df_raw, ajustadores_validos = sincronizar_y_cargar_datos(forzar_sync, dias_semana_target)

    # Pestañas de Navegación Gerencial
    tab_dashboard, tab_operacional, tab_gantt = st.tabs([
        "📈 Dashboard Ejecutivo (BI)", 
        "📋 Reporte Operacional de Equipo", 
        "📊 Carta Gantt Operativa"
    ])
    
    with tab_dashboard:
        renderizar_dashboard_ejecutivo(df_week, target_week_id, week_id_obj)
        
    with tab_operacional:
        renderizar_reporte_operacional(df_week, ajustadores_validos, target_week_id, week_id_obj)
        
    with tab_gantt:
        renderizar_carta_gantt(df_week, df_raw, dias_semana_target, target_week_id, week_id_obj)# ---------------------------------------------------------
# BLOQUE 4.0: ORQUESTADOR PRINCIPAL DE LA VISTA
# VERSIÓN: 4.0.2 (Blindaje de identificación semanal)
# ---------------------------------------------------------
def vista_reportes():
    import streamlit as st
    from datetime import datetime, timedelta
    
    st.title("📊 Tablero de Control y Planificación")
    
    col_radio, col_btn = st.columns([2, 1])
    with col_radio:
        week_id_obj = st.radio(
            "Seleccione la semana a reportar:", 
            ["Semana Pasada", "Semana Actual", "Próxima Semana"], 
            index=1,
            horizontal=True
        )
    with col_btn:
        st.markdown("<br>", unsafe_allow_html=True)
        forzar_sync = st.button("🔄 Sincronizar Nube ahora", type="primary", use_container_width=True)

    if week_id_obj == "Semana Pasada":
        offset = -1
    elif week_id_obj == "Semana Actual":
        offset = 0
    else:
        offset = 1
        
    hoy = datetime.now()
    # Calculamos la fecha base del periodo seleccionado
    fecha_referencia = hoy + timedelta(weeks=offset)
    # Obtenemos el lunes de esa semana
    lunes = fecha_referencia - timedelta(days=fecha_referencia.weekday())
    dias_semana_target = [(lunes + timedelta(days=i)).date() for i in range(7)]
        
    target_week_id = get_week_identifier(offset)

    # 4.2 Llamada al motor de datos
    df_week, df_raw, ajustadores_validos = sincronizar_y_cargar_datos(forzar_sync, dias_semana_target)

    tab_dashboard, tab_operacional, tab_gantt = st.tabs([
        "📈 Dashboard Ejecutivo (BI)", 
        "📋 Reporte Operacional de Equipo", 
        "📊 Carta Gantt Operativa"
    ])
    
    with tab_dashboard:
        renderizar_dashboard_ejecutivo(df_week, target_week_id, week_id_obj)
    with tab_operacional:
        renderizar_reporte_operacional(df_week, ajustadores_validos, target_week_id, week_id_obj)
    with tab_gantt:
        renderizar_carta_gantt(df_week, df_raw, dias_semana_target, target_week_id, week_id_obj)


# ---------------------------------------------------------
# BLOQUE PRINCIPAL: ENRUTADOR DE NAVEGACIÓN
# ---------------------------------------------------------
def main():
    st.sidebar.image("https://img.icons8.com/color/96/000000/engineering.png", width=60)
    st.sidebar.title("Navegación OpsControl")
    
    opcion = st.sidebar.radio(
        "Ir a:",
        ["Planificador Semanal", "Planificador Mensual MCL", "Programa Diario", "Reportes Jefatura"]
    )
    
    st.sidebar.markdown("---")
    
    if opcion == "Planificador Semanal":
        vista_planificador("Semanal")
    elif opcion == "Planificador Mensual MCL":
        vista_planificador("Mensual")
    elif opcion == "Programa Diario":
        vista_diario()
    elif opcion == "Reportes Jefatura":
        vista_reportes()

if __name__ == "__main__":
    main()
